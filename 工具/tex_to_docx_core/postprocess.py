from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

from .config import FONT_SIZE
from .docx_common import _apply_para, _heading_level
from .docx_structure import (
    _relocate_bibliography,
    add_heading_numbers,
    demote_heading4_to_heading3,
    insert_page_breaks_before_headings,
    insert_thesis_title_pages,
    normalize_special_h1_text,
)
from .front_matter import (
    compress_declaration_and_authorization,
    fill_cover_info,
    prepend_front_matter,
    tab_align_cover_and_signature_rows,
)
from .page_setup import _set_page_margins_a4, setup_page_numbers_and_sections
from .style import (
    bolden_abstract_prefixes,
    enable_latin_word_break,
    fold_abstract_heading_into_body,
    justify_body_paragraphs,
    normalize_bibliography_text,
    normalize_text_punctuation,
    normalize_all_fonts,
    normalize_paragraph_spacing,
    normalize_toc_entries,
    strip_cjk_latin_spaces,
)
from .tables import (
    add_equation_numbers,
    apply_table_body_font_size,
    apply_three_line_tables,
    center_all_images,
    center_all_table_cells,
    flatten_list_indent,
    resize_all_images_to_width,
    style_code_block_tables,
    wrap_listings_and_algorithms,
)

# =============================================================================
# Stage 3 — post-process via python-docx
# =============================================================================


_CODE_ALGO_CAPTION_RE = re.compile(r"^(?:代码|算法)\s*\d+[\.-]\d+\s")


def _migrate_heading_bold_to_style(doc) -> None:
    """把 Heading1/Heading2 的 "bold + 黑体" 从 run-level 迁移到 style-level。

    动机：Word 更新 TOC 域时，按不同实现版本，可能会把源 heading paragraph
    里的 run-level 格式（<w:b/>、<w:rFonts eastAsia="黑体"/>）拷贝到 TOC 项里，
    覆盖 TOC 样式自身的粗体定义。现在的规范是 TOC1 保持加粗黑体，TOC2 保持宋体
    不加粗，所以这里要把 heading 的粗体收回到 style-level，避免二级目录被污染。

    修复路径：
    1. 在 Heading1 / Heading2 style 的 rPr 上补 <w:b/> 和 eastAsia="黑体"，
       保证正文 H1/H2 视觉仍为黑体加粗（靠段落样式继承）。
    2. 遍历所有 H1/H2 paragraph runs，移除 run-level <w:b/>、<w:bCs/> 以及
       rFonts 的 eastAsia="黑体" 属性；让 run 不再携带 bold 格式 override，
       这样 TOC2 不会被 heading 的直设格式带成加粗。

    这样 Word/WPS 在"更新目录域"时拷 run-level 属性也拷不到 bold。"""
    from docx.oxml import OxmlElement as _OxmlElement
    from docx.oxml.ns import qn as _qn

    styles_el = doc.styles.element
    for style_id in ("Heading1", "Heading2", "Heading3", "Heading4", "Heading5", "Heading6"):
        style = None
        for s in styles_el.findall(_qn("w:style")):
            if s.get(_qn("w:styleId")) == style_id:
                style = s
                break
        if style is None:
            continue
        rPr = style.find(_qn("w:rPr"))
        if rPr is None:
            rPr = _OxmlElement("w:rPr")
            style.append(rPr)
        # 确保 rFonts 上 eastAsia=黑体（中文），ascii/hAnsi 继承 Times New Roman
        rFonts = rPr.find(_qn("w:rFonts"))
        if rFonts is None:
            rFonts = _OxmlElement("w:rFonts")
            rPr.insert(0, rFonts)
        # H1/H2 用黑体加粗，H3+ 用宋体加粗（与 _apply_para 保持一致）
        rFonts.set(_qn("w:eastAsia"), "黑体" if style_id in ("Heading1", "Heading2") else "宋体")
        rFonts.set(_qn("w:ascii"), "Times New Roman")
        rFonts.set(_qn("w:hAnsi"), "Times New Roman")
        rFonts.set(_qn("w:cs"), "Times New Roman")
        # 清除 theme-based font 属性，防止 theme 覆盖具体值
        for attr in ("w:asciiTheme", "w:hAnsiTheme", "w:eastAsiaTheme", "w:cstheme"):
            if rFonts.get(_qn(attr)) is not None:
                del rFonts.attrib[_qn(attr)]
        # 确保 <w:b/> 和 <w:bCs/> 存在
        for tag in ("w:b", "w:bCs"):
            if rPr.find(_qn(tag)) is None:
                rPr.append(_OxmlElement(tag))
        # 颜色强制黑色（覆盖模板 themeColor accent1 的深蓝；也让 heading 里嵌入的公式为黑）
        color = rPr.find(_qn("w:color"))
        if color is None:
            color = _OxmlElement("w:color")
            rPr.append(color)
        color.set(_qn("w:val"), "000000")
        for attr in ("w:themeColor", "w:themeShade", "w:themeTint"):
            if color.get(_qn(attr)) is not None:
                del color.attrib[_qn(attr)]

    # 遍历所有 Heading N paragraph runs，剥除 run-level bold 和 eastAsia=黑体
    for p in doc.paragraphs:
        sname = p.style.name
        if not sname.startswith("Heading "):
            continue
        for r in p.runs:
            rPr = r._element.find(_qn("w:rPr"))
            if rPr is None:
                continue
            for tag in ("w:b", "w:bCs"):
                for old in rPr.findall(_qn(tag)):
                    rPr.remove(old)
            rFonts = rPr.find(_qn("w:rFonts"))
            if rFonts is not None:
                # 让 eastAsia 由 style 接管；ascii/hAnsi 保持 Times New Roman
                if rFonts.get(_qn("w:eastAsia")) in ("黑体", "宋体"):
                    del rFonts.attrib[_qn("w:eastAsia")]


def _style_code_algorithm_captions(doc) -> None:
    """代码题 / 算法题标题段刷成表题规格（小四黑体加粗居中，段前6pt段后0）。
    flatten 阶段用 `\\noindent\\textbf{代码X-Y ...}` / `...{算法X.Y ...}` 生成整段粗体标题段，
    pandoc 不落 Caption 样式，主循环的表题分支覆盖不到，这里独立识别补刷。
    判定：段文字匹配 ^(代码|算法)N[-.]N 且段内所有非空 run 均粗体，过滤正文中"算法 5.1"这种局部加粗片段。
    """
    for p in doc.paragraphs:
        text = p.text.strip()
        if not _CODE_ALGO_CAPTION_RE.match(text):
            continue
        runs_with_text = [r for r in p.runs if (r.text or "").strip()]
        if not runs_with_text:
            continue
        if not all(r.bold for r in runs_with_text):
            continue
        _apply_para(
            p,
            cjk_font="黑体",
            ascii_font="Times New Roman",
            size_pt=FONT_SIZE["小四"],
            bold=True,
            align=WD_ALIGN_PARAGRAPH.CENTER,
            space_before=6.0,
            space_after=0.0,
        )


def _keep_captions_with_objects(doc) -> None:
    """给图片段、表题段、算法/代码题段加 keepNext，防止标题与主体跨页。"""
    from docx.oxml.ns import qn as _qn
    from docx.oxml import OxmlElement

    for i, p in enumerate(doc.paragraphs):
        style = p.style.name if p.style else ""
        text = p.text.strip()
        need_keep = False

        # 图片段（CaptionedFigure）：图片和下方图题不分页
        if style == "Captioned Figure":
            need_keep = True
        # 表题段：表题和下方表格不分页
        elif style == "Table Caption":
            need_keep = True
        # 代码/算法题段
        elif _CODE_ALGO_CAPTION_RE.match(text):
            need_keep = True

        if need_keep:
            pPr = p._element.find(_qn("w:pPr"))
            if pPr is None:
                pPr = OxmlElement("w:pPr")
                p._element.insert(0, pPr)
            if pPr.find(_qn("w:keepNext")) is None:
                pPr.append(OxmlElement("w:keepNext"))


def _unmono_code_styles(doc) -> None:
    """去掉 pandoc 给 inline code / code block 施加的等宽字体。

    pandoc 转 LaTeX `\\texttt{...}` 和 `\\begin{lstlisting}...` 时分别产出
    ``VerbatimChar`` 字符样式（Consolas）和 ``SourceCode`` 段落样式（link 到
    VerbatimChar）。用户规范要求全文统一宋体正文 + Times New Roman 西文，
    驼峰词 ``PathBoundsDecider`` 等 inline code 不要以 Consolas 等宽显示，
    代码块段落也统一回正文字形。

    做法：覆盖两个样式的 rFonts（ascii/hAnsi/cs = Times New Roman，eastAsia =
    宋体），清空 sz 让它继承段落大小。"""
    from docx.oxml import OxmlElement as _OxmlElement
    from docx.oxml.ns import qn as _qn

    styles_el = doc.styles.element
    targets = {"VerbatimChar", "SourceCode", "Verbatim"}
    for s in styles_el.findall(_qn("w:style")):
        sid = s.get(_qn("w:styleId"))
        if sid not in targets:
            continue
        rPr = s.find(_qn("w:rPr"))
        if rPr is None:
            rPr = _OxmlElement("w:rPr")
            s.append(rPr)
        rFonts = rPr.find(_qn("w:rFonts"))
        if rFonts is None:
            rFonts = _OxmlElement("w:rFonts")
            rPr.insert(0, rFonts)
        rFonts.set(_qn("w:ascii"), "Times New Roman")
        rFonts.set(_qn("w:hAnsi"), "Times New Roman")
        rFonts.set(_qn("w:cs"), "Times New Roman")
        rFonts.set(_qn("w:eastAsia"), "宋体")
        for attr in ("w:asciiTheme", "w:hAnsiTheme", "w:eastAsiaTheme", "w:cstheme"):
            if rFonts.get(_qn(attr)) is not None:
                del rFonts.attrib[_qn(attr)]
        # 清掉 sz，让段落/上下文决定字号
        for tag in ("w:sz", "w:szCs"):
            for old in rPr.findall(_qn(tag)):
                rPr.remove(old)

    # run-level 清洗：正文里若直接给 run 设了 Consolas/Courier（docx 有时写死），改为 Times New Roman
    mono_names = {"Consolas", "Courier New", "Courier", "DejaVu Sans Mono", "Monaco", "Menlo"}
    for p in doc.paragraphs:
        for r in p.runs:
            rPr = r._element.find(_qn("w:rPr"))
            if rPr is None:
                continue
            rFonts = rPr.find(_qn("w:rFonts"))
            if rFonts is None:
                continue
            changed = False
            for attr in ("w:ascii", "w:hAnsi", "w:cs"):
                val = rFonts.get(_qn(attr))
                if val in mono_names:
                    rFonts.set(_qn(attr), "Times New Roman")
                    changed = True
            if rFonts.get(_qn("w:eastAsia")) in mono_names:
                rFonts.set(_qn("w:eastAsia"), "宋体")
                changed = True
            _ = changed  # 显式标记一下方便调试

    # 表格 cell 里的 runs 同样处理（代码块被包进单格表）
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for r in p.runs:
                        rPr = r._element.find(_qn("w:rPr"))
                        if rPr is None:
                            continue
                        rFonts = rPr.find(_qn("w:rFonts"))
                        if rFonts is None:
                            continue
                        for attr in ("w:ascii", "w:hAnsi", "w:cs"):
                            val = rFonts.get(_qn(attr))
                            if val in mono_names:
                                rFonts.set(_qn(attr), "Times New Roman")
                        if rFonts.get(_qn("w:eastAsia")) in mono_names:
                            rFonts.set(_qn("w:eastAsia"), "宋体")


def _force_hyperlinks_black(doc) -> None:
    """论文规范要求全文黑色，包括参考文献里的 DOI/URL 超链接。
    pandoc 默认超链接是 Word 的 Hyperlink 字符样式（蓝+下划线）；
    这里对所有 ``<w:hyperlink>`` 内部 run 的 ``<w:color>`` 强制设为 000000，
    同时把字符样式从 Hyperlink 解绑，避免 Word 重置颜色。"""
    from docx.oxml import OxmlElement as _OxmlElement
    from docx.oxml.ns import qn as _qn

    hyps = doc.element.body.findall('.//' + _qn('w:hyperlink'))
    for h in hyps:
        for r in h.findall('.//' + _qn('w:r')):
            rPr = r.find(_qn('w:rPr'))
            if rPr is None:
                rPr = _OxmlElement('w:rPr')
                r.insert(0, rPr)
            # 解绑 Hyperlink 字符样式，让颜色设定不被样式重置
            for rStyle in list(rPr.findall(_qn('w:rStyle'))):
                if rStyle.get(_qn('w:val')) in ('Hyperlink', 'FootnoteReference'):
                    rPr.remove(rStyle)
            # 强制 color=000000
            color = rPr.find(_qn('w:color'))
            if color is None:
                color = _OxmlElement('w:color')
                rPr.append(color)
            color.set(_qn('w:val'), '000000')
            # 去掉下划线（themeColor 可能也带蓝），保留纯黑文本
            for u in list(rPr.findall(_qn('w:u'))):
                rPr.remove(u)


def post_process(docx_path: Path) -> None:
    """docx 后处理总入口，按 6 个 Stage 顺序执行。

    Stage A：章节元信息（anchor 迁移 / Heading 降级 / 章节编号 / 分页 / 摘要空行）
    Stage B：主样式循环（按 Heading/Caption 落 cjk_font / ascii_font / bold / spacing）
    Stage C：前置页 + 版面（封面/声明/授权/目录 + 页边距页码 + 三线表 + 图片 + 公式 + 代码块）
    Stage D：文本 + 字符级（列表编号 / 签名页 tab / 半角引号 / 参考文献细节）
    Stage E：全局规格化（字体 / 对齐 / 行距 / 超链接 / Heading bold 迁移）
    Stage F：规则驱动闭环（_normalize_for_inspector 兜底层 + _apply_format_rules 规则层）
    """
    doc = Document(str(docx_path))

    # ---------------------------------------------------------------
    # Stage A：章节元信息
    # ---------------------------------------------------------------
    _relocate_bibliography(doc)                 # A1 参考文献条目搬到 anchor 标题之后
    demote_heading4_to_heading3(doc)            # A2 Heading 4 → Heading 3（规范"最多三级标题"）
    normalize_special_h1_text(doc)              # A3 致谢 → 致 谢（规范文字要求）
    add_heading_numbers(doc)                    # A4 给 Heading 1/2/3 加章节编号前缀
    insert_page_breaks_before_headings(doc)     # A5 所有 Heading 1 前插分页符（首个除外）
    _insert_blank_before_abstract(doc)          # A6 摘要/Abstract 段前空行

    # ---------------------------------------------------------------
    # Stage B：主样式循环（按段类型落字体/字号/加粗/spacing）
    # ---------------------------------------------------------------
    for p in doc.paragraphs:
        style = p.style.name if p.style else ""
        text = p.text.strip()
        lvl = _heading_level(style)

        if lvl == 1:
            # 致谢/参考文献/成果/附录：三号黑体加粗居中（规范要求）
            if (
                text in ("致 谢", "致谢", "参考文献", "本科期间的学习与科研成果")
                or text.startswith("附录")
            ):
                _apply_para(
                    p,
                    cjk_font="黑体",
                    ascii_font="Times New Roman",
                    size_pt=FONT_SIZE["三号"],
                    bold=True,
                    align=WD_ALIGN_PARAGRAPH.CENTER,
                    space_before=12.0,
                    space_after=6.0,
                )
            else:
                # 其它章 / 摘要 / Abstract — 四号黑体加粗左对齐
                _apply_para(
                    p,
                    cjk_font="黑体",
                    ascii_font="Times New Roman",
                    size_pt=FONT_SIZE["四号"],
                    bold=True,
                    align=WD_ALIGN_PARAGRAPH.LEFT,
                    space_before=12.0,
                    space_after=6.0,
                )
        elif lvl == 2:
            _apply_para(
                p,
                cjk_font="黑体",
                ascii_font="Times New Roman",
                size_pt=FONT_SIZE["小四"],
                bold=True,
                align=WD_ALIGN_PARAGRAPH.LEFT,
                space_before=6.0,
                space_after=3.0,
            )
        elif lvl == 3 or lvl == 4:
            # Heading 4 视觉上下沉为三级标题样式（规范"最多三级"）
            _apply_para(
                p,
                cjk_font="宋体",
                ascii_font="Times New Roman",
                size_pt=FONT_SIZE["小四"],
                bold=True,
                align=WD_ALIGN_PARAGRAPH.LEFT,
                space_before=3.0,
                space_after=3.0,
            )
        elif style in ("Caption", "Image Caption", "Table Caption"):
            # 图题/表题：小四黑体加粗居中（用户要求黑体），显式清除 Caption 样式的 italic 继承
            # 区分图题 vs 表题：按段落首字符"图"/"表"判定
            #   - 图题：段前 0、段后 6 磅（图在上、图题在下、与正文留白）
            #   - 表题：段前 6 磅、段后 0（表题在上、表在下、与正文留白）
            is_table_cap = text.startswith("表")
            _apply_para(
                p,
                cjk_font="黑体",
                ascii_font="Times New Roman",
                size_pt=FONT_SIZE["小四"],
                bold=True,
                align=WD_ALIGN_PARAGRAPH.CENTER,
                space_before=6.0 if is_table_cap else 0.0,
                space_after=0.0 if is_table_cap else 6.0,
            )
            for run in p.runs:
                run.font.italic = False  # 根除 pandoc Caption 样式默认 italic
        elif style == "Captioned Figure":
            # 图本体段，居中；不覆盖图片本身的 run
            pf = p.paragraph_format
            pf.line_spacing = 1.5
            pf.space_before = Pt(3.0)
            pf.space_after = Pt(3.0)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif style.lower().startswith(("bibliograph", "reference")):
            # pandoc 参考文献：中文宋体小四 / 英文 TNR 小四 / 固定行距 20 磅（规范 R5）
            _apply_para(
                p,
                cjk_font="宋体",
                ascii_font="Times New Roman",
                size_pt=FONT_SIZE["小四"],
                bold=False,
                align=WD_ALIGN_PARAGRAPH.JUSTIFY,
            )
            # 覆盖 line_spacing 为 固定 20 磅
            from docx.enum.text import WD_LINE_SPACING

            pf = p.paragraph_format
            pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
            pf.line_spacing = Pt(20)
        else:
            # 正文
            _apply_para(
                p,
                cjk_font="宋体",
                ascii_font="Times New Roman",
                size_pt=FONT_SIZE["小四"],
                bold=False,
                first_line_chars=2 if text else 0,
            )

    # 表格单元格也统一样式
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    _apply_para(
                        p,
                        cjk_font="宋体",
                        ascii_font="Times New Roman",
                        size_pt=FONT_SIZE["小四"],
                        bold=False,
                        align=WD_ALIGN_PARAGRAPH.CENTER,
                        line_spacing=1.0,
                    )

    # ---------------------------------------------------------------
    # Stage C：前置页 + 版面（封面 / 声明 / 授权 / 目录 / 页码 / 图表 / 代码）
    # ---------------------------------------------------------------
    insert_thesis_title_pages(doc)              # C1 中英文论文题目页（避开 Normal 覆盖字号）
    prepend_front_matter(doc)                   # C2 模板封面+声明+授权+目录搬到正文前
    fill_cover_info(doc)                        # C3 封面信息填充（题目/学院/姓名/...）
    _set_page_margins_a4(doc)                   # C4 A4 页边距
    setup_page_numbers_and_sections(doc)        #    + 分节 / 页码 / 页脚
    apply_three_line_tables(doc)                # C5 表格三线化 + 整表居中
    center_all_images(doc)                      #    + 图片段居中
    resize_all_images_to_width(doc, width_cm=14.7, max_height_cm=9.9)  # + 图片等比缩放
    add_equation_numbers(doc)                   # C6 独立公式右对齐 "（N-M）"
    wrap_listings_and_algorithms(doc)           # C7 代码清单+算法伪代码 → 单格全外框表（必须在三线表后）
    style_code_block_tables(doc)                #    + 代码块表格：五号/不缩进/左对齐
    center_all_table_cells(doc)                 #    + 表格 cell 居中（代码块只垂直居中）
    apply_table_body_font_size(doc)             #    + 数据型表格 cell 字号五号
    _style_code_algorithm_captions(doc)         # C8 代码/算法题标题段刷成表题规格
    _keep_captions_with_objects(doc)             # C9 图/表/算法题与主体不跨页

    # ---------------------------------------------------------------
    # Stage D：文本 + 字符级
    # ---------------------------------------------------------------
    flatten_list_indent(doc)                    # D1 列表自动编号顶格
    compress_declaration_and_authorization(doc) # D2 压缩声明+授权同页
    tab_align_cover_and_signature_rows(doc)     # D3 封面/签名页"左标签+右字段"tab 对齐
    fold_abstract_heading_into_body(doc)        # D4 折叠 H1 "摘要"/"Abstract" 标题段
    bolden_abstract_prefixes(doc)               # D5 摘要/关键词前缀黑体加粗（必须在主循环后）
    normalize_text_punctuation(doc)             # D6 半角引号→全角（避开代码块）
    normalize_bibliography_text(doc)            # D7 参考文献英文小圆点后补空格

    # ---------------------------------------------------------------
    # Stage E：全局规格化
    # ---------------------------------------------------------------
    normalize_all_fonts(doc)                    # E1 字体正则化（宋体+TNR，CJK_SAFE 含黑体不覆盖前缀）
    justify_body_paragraphs(doc)                # E2 正文段两端对齐（所有段已定型后）
    normalize_paragraph_spacing(doc)            # E3 段间距/首行缩进/snapToGrid
    enable_latin_word_break(doc)                # E4 允许长驼峰词（PathBoundsDecider）中间拆行
    normalize_toc_entries(doc)                  # E5 目录项缩进/段后收口
    _force_hyperlinks_black(doc)                # E6 超链接（DOI/URL）强制黑色
    _migrate_heading_bold_to_style(doc)         # E7 H1/H2 bold+黑体 run→style（防 TOC 域污染）
    _unmono_code_styles(doc)                    # E8 去 pandoc 给 \texttt/代码块的 Consolas

    # ---------------------------------------------------------------
    # Stage F：规则驱动闭环（FORMAT_RULES.md）
    # ---------------------------------------------------------------
    _normalize_for_inspector(doc)               # F1 兜底层：caption 前缀补回、数学对象字号、TOC tab leader
    _close_inspector_issues(doc)                # F2 规则层：_apply_format_rules 按 7 块规则链

    # ---------------------------------------------------------------
    # Stage G：最终文本清理（必须在 F1 的空格插入之后）
    # ---------------------------------------------------------------
    strip_cjk_latin_spaces(doc)                 # G1 删除中英文间手动空格（Word autoSpace 管理）

    doc.save(str(docx_path))

    # 历史注记：早期曾做 PDF-driven 静态化（toc_freeze 模块）规避 WPS 对
    # hyperlink 嵌套 PAGEREF 的解析 bug（fallback 全 6）。现根因已在
    # `_make_toc_entry_paragraph` 解决——PAGEREF 域拆出 hyperlink 外部，配合
    # settings.updateFields=true 由 Word/WPS 自己更新域。静态化模块已移除。


_INSPECTOR_FIG_CAP_RE = re.compile(r"^\s*图\s*\d+[\.\-]\d+")
_INSPECTOR_TAB_CAP_RE = re.compile(r"^\s*表\s*\d+[\.\-]\d+")
_INSPECTOR_CODE_CAP_RE = re.compile(r"^\s*(代码|算法)\s*\d+[\.\-]\d+")


def _normalize_for_inspector(doc) -> None:
    """按学校格式检测器口径做最后统一修正，闭环修复检测报告列出的错误。

    规则来源——thesis_格式检测报告.html：
    - 图 caption：段前 0 段后 0（当前 after=120 违规）。
    - 表 caption：段前 0 段后 0（当前 before=120 违规）。
    - 代码/算法 caption：宋体常规，居中，无首行缩进，段前 0 段后 0（当前是黑体加粗 before=120）。
    - 摘要/Abstract 首段：段前 0 段后 0（当前 before=120 after=31 违规）。
    - 摘要/Abstract/目录 题目段：行距 1.5 倍（当前单倍）。
    - "目录"说明段："（打开文档后…）"行距 1.5。
    - 成果页标题 "本科期间的学习与科研成果"：居中对齐。
    - 所有 caption 段 run 的 eastAsia 强制宋体（当前部分 caption 混黑体）。
    """
    from docx.oxml import OxmlElement as _OxmlElement
    from docx.oxml.ns import qn as _qn

    WNS = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"

    def _ensure(child, tag):
        el = child.find(_qn(tag))
        if el is None:
            el = _OxmlElement(tag)
            child.append(el)
        return el

    def _set_zero_spacing(pPr):
        sp = _ensure(pPr, "w:spacing")
        sp.set(_qn("w:before"), "0")
        sp.set(_qn("w:after"), "0")
        sp.set(_qn("w:beforeLines"), "0")
        sp.set(_qn("w:afterLines"), "0")
        sp.set(_qn("w:beforeAutospacing"), "0")
        sp.set(_qn("w:afterAutospacing"), "0")

    def _set_line_1p5(pPr):
        sp = _ensure(pPr, "w:spacing")
        sp.set(_qn("w:line"), "360")
        sp.set(_qn("w:lineRule"), "auto")

    def _clear_firstline(pPr):
        ind = pPr.find(_qn("w:ind"))
        if ind is None:
            return
        for attr in ("firstLine", "firstLineChars"):
            k = _qn(f"w:{attr}")
            if k in ind.attrib:
                del ind.attrib[k]

    def _force_run_songti(r_el, *, unbold=False, bold=False):
        rPr = r_el.find(_qn("w:rPr"))
        if rPr is None:
            rPr = _OxmlElement("w:rPr")
            r_el.insert(0, rPr)
        rF = rPr.find(_qn("w:rFonts"))
        if rF is None:
            rF = _OxmlElement("w:rFonts")
            rPr.insert(0, rF)
        rF.set(_qn("w:eastAsia"), "宋体")
        if not rF.get(_qn("w:ascii")):
            rF.set(_qn("w:ascii"), "Times New Roman")
        if not rF.get(_qn("w:hAnsi")):
            rF.set(_qn("w:hAnsi"), "Times New Roman")
        if not rF.get(_qn("w:cs")):
            rF.set(_qn("w:cs"), "Times New Roman")
        for attr in ("w:asciiTheme", "w:hAnsiTheme", "w:eastAsiaTheme", "w:cstheme"):
            k = _qn(attr)
            if k in rF.attrib:
                del rF.attrib[k]
        if unbold:
            # 显式关闭加粗（覆盖父级样式/run 级）
            for tag in ("w:b", "w:bCs"):
                for old in list(rPr.findall(_qn(tag))):
                    rPr.remove(old)
                neg = _OxmlElement(tag)
                neg.set(_qn("w:val"), "0")
                rPr.append(neg)
        elif bold:
            # 显式开启加粗（caption 家族要求加粗）
            for tag in ("w:b", "w:bCs"):
                for old in list(rPr.findall(_qn(tag))):
                    rPr.remove(old)
                pos = _OxmlElement(tag)
                # 不设 val 等价于 val=true
                rPr.append(pos)

    def _para_text(p_el):
        return "".join((t.text or "") for t in p_el.findall(".//" + _qn("w:t")))

    def _para_style(p_el):
        pPr = p_el.find(_qn("w:pPr"))
        if pPr is None:
            return ""
        pStyle = pPr.find(_qn("w:pStyle"))
        if pStyle is None:
            return ""
        return pStyle.get(_qn("w:val")) or ""

    body = doc.element.body

    # —— Pass 0：给 pandoc 跨 OMML 吞掉前缀的 ImageCaption / TableCaption 段补回 "图 N.M " / "表 N.M "。
    # 状态机：按 body 直属顺序扫描，H1 递增章号并清零图/表计数，ImageCaption / TableCaption 段内递增。
    # 若段文本已带 "图 N.M" / "表 N.M" 前缀则跳过；仅对 pandoc 因 caption 内含 OMML 导致前缀被吞的段补回。
    _chap = 0
    _fig_n = 0
    _tab_n = 0
    for child in list(body):
        if child.tag != _qn("w:p"):
            continue
        style_val = _para_style(child)
        if style_val == "Heading1":
            _chap += 1
            _fig_n = 0
            _tab_n = 0
            continue
        if style_val not in ("ImageCaption", "TableCaption"):
            continue
        txt = _para_text(child).strip()
        if not txt:
            continue
        if style_val == "ImageCaption":
            _fig_n += 1
            if _INSPECTOR_FIG_CAP_RE.match(txt):
                continue  # 已带前缀
            prefix_text = f"图{_chap}.{_fig_n} "
        else:
            _tab_n += 1
            if _INSPECTOR_TAB_CAP_RE.match(txt):
                continue
            prefix_text = f"表{_chap}.{_tab_n} "
        # 构造前缀 run：宋体加粗小四，与 caption 其他 run 样式对齐
        new_r = _OxmlElement("w:r")
        new_rPr = _OxmlElement("w:rPr")
        rF = _OxmlElement("w:rFonts")
        rF.set(_qn("w:ascii"), "Times New Roman")
        rF.set(_qn("w:hAnsi"), "Times New Roman")
        rF.set(_qn("w:cs"), "Times New Roman")
        rF.set(_qn("w:eastAsia"), "宋体")
        new_rPr.append(rF)
        new_rPr.append(_OxmlElement("w:b"))
        new_rPr.append(_OxmlElement("w:bCs"))
        sz = _OxmlElement("w:sz")
        sz.set(_qn("w:val"), "24")
        new_rPr.append(sz)
        szCs = _OxmlElement("w:szCs")
        szCs.set(_qn("w:val"), "24")
        new_rPr.append(szCs)
        new_r.append(new_rPr)
        new_t = _OxmlElement("w:t")
        new_t.set(_qn("xml:space"), "preserve")
        new_t.text = prefix_text
        new_r.append(new_t)
        # 找第一个内容节点（w:r / m:oMath / m:oMathPara），插在它之前
        _M = "{http://schemas.openxmlformats.org/officeDocument/2006/math}"
        content_tags = {_qn("w:r"), f"{_M}oMath", f"{_M}oMathPara"}
        first_content = None
        for c in child:
            if c.tag in content_tags:
                first_content = c
                break
        if first_content is not None:
            first_content.addprevious(new_r)
        else:
            child.append(new_r)

    for child in list(body):
        if child.tag != _qn("w:p"):
            continue
        txt = _para_text(child).strip()
        if not txt:
            continue
        style_val = _para_style(child)
        pPr = child.find(_qn("w:pPr"))
        if pPr is None:
            pPr = _OxmlElement("w:pPr")
            child.insert(0, pPr)

        is_img_cap = (
            style_val == "ImageCaption"
            or (style_val == "Caption" and txt.startswith("图"))
            or (_INSPECTOR_FIG_CAP_RE.match(txt) and len(txt) < 120)
        )
        is_tab_cap = (
            style_val == "TableCaption"
            or (style_val == "Caption" and txt.startswith("表"))
            or (_INSPECTOR_TAB_CAP_RE.match(txt) and len(txt) < 120)
        )
        is_code_cap = bool(_INSPECTOR_CODE_CAP_RE.match(txt)) and len(txt) < 80

        if is_code_cap:
            _set_zero_spacing(pPr)
            _clear_firstline(pPr)
            jc = _ensure(pPr, "w:jc")
            jc.set(_qn("w:val"), "center")
            for r_el in child.findall(_qn("w:r")):
                _force_run_songti(r_el, bold=True)
            continue

        if is_img_cap or is_tab_cap:
            _set_zero_spacing(pPr)
            _clear_firstline(pPr)
            for r_el in child.findall(_qn("w:r")):
                _force_run_songti(r_el, bold=True)
            continue

        # 摘要/Abstract 首段：段前 0 段后 0
        txt_strip = txt.lstrip()
        if txt_strip.startswith(("摘要：", "摘要:", "Abstract:", "Abstract：")) or (
            txt_strip.startswith("Abstract") and len(txt_strip) > 8 and txt_strip[8] in (" ", ":", "：")
        ):
            _set_zero_spacing(pPr)
            continue

        # 成果页标题居中
        if txt == "本科期间的学习与科研成果":
            jc = _ensure(pPr, "w:jc")
            jc.set(_qn("w:val"), "center")
            continue

    # 摘要/Abstract/目录 的居中题目段 + 目录说明段：行距 1.5
    for child in list(body):
        if child.tag != _qn("w:p"):
            continue
        txt = _para_text(child).strip()
        pPr = child.find(_qn("w:pPr"))
        if pPr is None:
            continue
        jc = pPr.find(_qn("w:jc"))
        jc_val = jc.get(_qn("w:val")) if jc is not None else ""
        # 论文题目段（居中、非 caption、非目录）
        if jc_val == "center" and txt and not (
            _INSPECTOR_FIG_CAP_RE.match(txt)
            or _INSPECTOR_TAB_CAP_RE.match(txt)
            or _INSPECTOR_CODE_CAP_RE.match(txt)
            or txt in ("目  录", "目录")
        ):
            _set_line_1p5(pPr)
            continue
        # 目录/Abstract/摘要标题段
        if txt in ("目  录", "目录"):
            _set_line_1p5(pPr)
            continue
        # 目录说明段
        if txt.startswith("（打开文档"):
            _set_line_1p5(pPr)
            continue

    # 参考文献段：悬挂缩进按编号位数 (1 位数 175, 2 位及以上 200) + 1.5 倍行距
    for p in doc.paragraphs:
        style_name = (p.style.name if p.style else "").lower()
        if not style_name.startswith(("bibliograph", "reference")):
            continue
        pPr = p._element.get_or_add_pPr()
        ind = pPr.find(_qn("w:ind"))
        if ind is None:
            ind = _OxmlElement("w:ind")
            pPr.append(ind)
        # 清 firstLine，设 hangingChars 随编号宽度
        for attr in ("firstLine", "firstLineChars", "hanging"):
            k = _qn(f"w:{attr}")
            if k in ind.attrib:
                del ind.attrib[k]
        # 用正则抓条目编号：[1]~[9] → 175, [10]+ → 200
        ref_text = p.text or ""
        m = re.match(r"\s*\[(\d+)\]", ref_text)
        if m and len(m.group(1)) >= 2:
            hang_chars = "200"
        else:
            hang_chars = "175"
        ind.set(_qn("w:hangingChars"), hang_chars)
        ind.set(_qn("w:leftChars"), hang_chars)
        # 行距 1.5 倍（pandoc 默认 400/exact=固定20pt，学校要求 360/auto）
        sp = pPr.find(_qn("w:spacing"))
        if sp is None:
            sp = _OxmlElement("w:spacing")
            pPr.append(sp)
        sp.set(_qn("w:line"), "360")
        sp.set(_qn("w:lineRule"), "auto")

    # 中文段落英文半角逗号 → 全角（跨 run 映射）
    cjk_re = re.compile(r"[一-鿿]")
    half_comma_re = re.compile(r"[一-鿿]\s*,\s*[一-鿿]")

    def _has_cjk(s):
        return bool(cjk_re.search(s))

    for p in doc.paragraphs:
        style_name = (p.style.name if p.style else "")
        # 跳过代码块/verbatim、参考文献
        if style_name.startswith(("Source", "Verbatim", "Reference", "Bibliograph")):
            continue
        text = p.text or ""
        if not _has_cjk(text) or "," not in text:
            continue
        # 找所有半角逗号的段内绝对位置，要求其前一个 CJK 字符 + 后一个 CJK 字符
        comma_positions = []
        for m in re.finditer(r",", text):
            i = m.start()
            # 左侧最近非空字符
            li = i - 1
            while li >= 0 and text[li].isspace():
                li -= 1
            if li < 0 or not cjk_re.match(text[li]):
                continue
            # 右侧最近非空字符
            ri = i + 1
            while ri < len(text) and text[ri].isspace():
                ri += 1
            if ri >= len(text) or not cjk_re.match(text[ri]):
                continue
            comma_positions.append(i)
        if not comma_positions:
            continue
        # 把段内 run 文本的绝对区间算出来，逐 run 重写
        # 同时删除全角逗号后紧跟的空格（中文标点后不需要空格）
        offset = 0
        for r in p.runs:
            rt = r.text or ""
            r_start = offset
            r_end = offset + len(rt)
            # 落在本 run 的半角逗号位置
            hits = [cp - r_start for cp in comma_positions if r_start <= cp < r_end]
            if hits:
                chars = list(rt)
                for h in sorted(hits, reverse=True):
                    if 0 <= h < len(chars) and chars[h] == ",":
                        chars[h] = "，"
                        # 删除逗号后紧跟的空格
                        while h + 1 < len(chars) and chars[h + 1] == " ":
                            chars.pop(h + 1)
                r.text = "".join(chars)
            offset = r_end

    # 全角逗号/顿号后跨 run 空格清理
    for p in doc.paragraphs:
        runs = p.runs
        for i in range(len(runs) - 1):
            cur = runs[i].text or ""
            if cur and cur[-1] in "，、；：":
                nxt = runs[i + 1].text or ""
                if nxt and nxt[0] == " ":
                    runs[i + 1].text = nxt.lstrip(" ")

    # CJK ↔ ASCII 交界处空格管理：已改为由 Word autoSpaceDE/DN 自动处理，
    # strip_cjk_latin_spaces (Stage G) 负责删除手动空格。此处不再插入。

    # OMML 数学对象字号统一小四 24 半磅（12pt）
    M = "{http://schemas.openxmlformats.org/officeDocument/2006/math}"
    for mr in body.iter(f"{M}r"):
        rPr = mr.find(f"{M}rPr")
        # 在 OMML 里字号用 w:sz 挂在 w:rPr 下，而不是 m:rPr；要插一个 w:rPr
        w_rPr = None
        for child in mr:
            if child.tag == _qn("w:rPr"):
                w_rPr = child
                break
        if w_rPr is None:
            w_rPr = _OxmlElement("w:rPr")
            # 放在 m:rPr 之后、m:t 之前
            insert_at = 0
            for idx, ch in enumerate(mr):
                if ch.tag == f"{M}rPr":
                    insert_at = idx + 1
                    break
            mr.insert(insert_at, w_rPr)
        sz = w_rPr.find(_qn("w:sz"))
        if sz is None:
            sz = _OxmlElement("w:sz")
            w_rPr.append(sz)
        sz.set(_qn("w:val"), "24")
        szCs = w_rPr.find(_qn("w:szCs"))
        if szCs is None:
            szCs = _OxmlElement("w:szCs")
            w_rPr.append(szCs)
        szCs.set(_qn("w:val"), "24")

    # 目录段：添加前导符（tab leader = dot）
    for p in doc.paragraphs:
        t = (p.text or "").strip()
        style_name = (p.style.name if p.style else "").lower()
        if "toc" in style_name or "目录" in style_name or style_name.startswith("contents"):
            pPr = p._element.find(_qn("w:pPr"))
            if pPr is None:
                continue
            tabs = pPr.find(_qn("w:tabs"))
            if tabs is None:
                tabs = _OxmlElement("w:tabs")
                pPr.append(tabs)
            # 确保有一个右对齐 tab 带 dot leader
            has_right_dot = False
            for tab in tabs.findall(_qn("w:tab")):
                if tab.get(_qn("w:val")) == "right":
                    tab.set(_qn("w:leader"), "dot")
                    has_right_dot = True
            if not has_right_dot:
                tab_el = _OxmlElement("w:tab")
                tab_el.set(_qn("w:val"), "right")
                tab_el.set(_qn("w:leader"), "dot")
                tab_el.set(_qn("w:pos"), "8306")
                tabs.append(tab_el)


def _insert_blank_before_abstract(doc) -> None:
    """在"摘要："/"Abstract:"所在段之前各插入一个空段，让论文题目与摘要正文之间留一空行过渡。
    幂等：若紧前一段已经是空段则跳过。"""
    from docx.oxml.ns import qn as _qn
    from docx.oxml import OxmlElement as _OxmlElement

    prefixes = ("摘要：", "摘要:", "Abstract:", "Abstract：")
    body = doc.element.body
    children = list(body)
    inserted = set()
    for idx, child in enumerate(children):
        if child.tag != _qn("w:p"):
            continue
        txt = "".join(
            (t.text or "") for t in child.findall(".//" + _qn("w:t"))
        ).lstrip()
        matched = next((pfx for pfx in prefixes if txt.startswith(pfx)), None)
        if matched is None:
            continue
        kind = "zh" if matched.startswith("摘要") else "en"
        if kind in inserted:
            continue
        if idx > 0:
            prev = children[idx - 1]
            if prev.tag == _qn("w:p"):
                prev_txt = "".join(
                    (t.text or "") for t in prev.findall(".//" + _qn("w:t"))
                )
                if not prev_txt.strip():
                    inserted.add(kind)
                    continue
        blank = _OxmlElement("w:p")
        child.addprevious(blank)
        inserted.add(kind)




# =============================================================================
# Stage 4 — 按 FORMAT_RULES.md 规则驱动的统一格式处理器
# =============================================================================
#
# 设计原则：
#   1. 先按文档章节（封面/目录/摘要/正文/参考文献/致谢/成果）切块
#   2. 每一块有自己的规则集（FORMAT_RULES.md §2）
#   3. 跨块的字符级规则（字体、半角逗号、caption 前缀空格等）最后统一跑
#   4. 不再按"检测器新错误"叠加补丁；新规则加到 FORMAT_RULES.md 对应章节即可
# =============================================================================

_CAP_LEAD_FIG_RE = re.compile(r"^(图)([\s  　]+)(\d)")
_CAP_LEAD_TAB_RE = re.compile(r"^(表)([\s  　]+)(\d)")
_CAP_LEAD_CODE_RE = re.compile(r"^(代码|算法)([\s  　]+)(\d)")
_CAP_ANY_RE = re.compile(r"^(图|表|代码|算法)\s*\d+[\.\-]\d+")
_CODE_ALGO_SHORT_CAP_RE = re.compile(r"^(代码|算法)\s*\d+[\.-]\d+\s+\S")
_EQN_NUMBER_RE = re.compile(r"^\s*[（(]\s*\d+[\.\-]\d+\s*[）)]\s*$")
_CJK_COMMA_RE = re.compile(r"([一-鿿])\s*,\s*([一-鿿])")
_BRACKET_REF_RE = re.compile(r"^\s*\[\d+\]\s")
_ABSTRACT_ZH_PFX = ("摘要：", "摘要:")
_ABSTRACT_EN_PFX = ("Abstract:", "Abstract：")


# ---------- 共用工具（从 docx_common 导入，本地只留 shortcut） ----------


def _qn_(tag):  # shortcut
    from docx.oxml.ns import qn as _qn
    return _qn(tag)


def _el_(tag):  # shortcut
    from docx.oxml import OxmlElement as _OxmlElement
    return _OxmlElement(tag)


# 以下通用工具统一从 docx_common 提供，postprocess 只做 re-export 以保持旧调用兼容
from .docx_common import (  # noqa: E402
    _all_runs_bold,
    _ensure_child,
    _ensure_pPr,
    _pstyle,
    _ptext,
    _run_clear_bold,
    _run_is_bold,
    _run_set_bold,
    _set_indent,
    _set_jc,
    _set_spacing,
)


# ---------- §1 块划分 ----------


def _identify_blocks(p_children):
    """按顺序扫描 p_children，识别章节块的起止索引。

    返回 dict: {block_name: (start, end)}  end 是 exclusive。
    规则来源：FORMAT_RULES.md §1。
    """
    blocks = {}
    n = len(p_children)
    # 先找关键锚点索引
    idx_toc_title = None          # "目  录" 段
    idx_zh_title = None            # 中文论文题目段（摘要页首行，居中 16pt）
    idx_en_title = None            # 英文论文题目段
    idx_zh_abs_body = None         # "摘要：" 段
    idx_en_abs_body = None         # "Abstract:" 段
    idx_h1_first = None            # 正文第一个 H1（"1 绪论"）
    idx_biblio_h1 = None
    idx_ack_h1 = None
    idx_achieve_h1 = None

    for i, p in enumerate(p_children):
        txt = _ptext(p).strip()
        style = _pstyle(p)
        is_h1 = style == "Heading1"
        is_center_16 = False
        pPr_c = p.find(_qn_("w:pPr"))
        if pPr_c is not None:
            jc_c = pPr_c.find(_qn_("w:jc"))
            if jc_c is not None and jc_c.get(_qn_("w:val")) == "center":
                for r_el in p.findall(_qn_("w:r")):
                    rPr_r = r_el.find(_qn_("w:rPr"))
                    if rPr_r is None:
                        continue
                    sz = rPr_r.find(_qn_("w:sz"))
                    if sz is not None and sz.get(_qn_("w:val")) == "32":
                        is_center_16 = True
                        break
        if idx_toc_title is None and txt in ("目录", "目 录", "目  录"):
            idx_toc_title = i
        if idx_zh_abs_body is None and txt.startswith(_ABSTRACT_ZH_PFX):
            idx_zh_abs_body = i
        if idx_en_abs_body is None and (
            txt.startswith(_ABSTRACT_EN_PFX)
            or (txt.startswith("Abstract") and len(txt) > 8 and txt[8] in " :：")
        ):
            idx_en_abs_body = i
        if idx_h1_first is None and is_h1 and (txt.startswith("1 ") or txt.startswith("1绪论")):
            idx_h1_first = i
        if idx_biblio_h1 is None and is_h1 and "参考文献" in txt:
            idx_biblio_h1 = i
        if idx_ack_h1 is None and is_h1 and ("致谢" in txt or "致 谢" in txt):
            idx_ack_h1 = i
        if idx_achieve_h1 is None and is_h1 and "本科期间的学习与科研成果" in txt:
            idx_achieve_h1 = i

    # 中文题目 = zh_abs_body 前紧邻的居中 16pt 段；英文题目 = en_abs_body 前紧邻的居中 16pt 段
    if idx_zh_abs_body is not None:
        for back in range(idx_zh_abs_body - 1, max(-1, idx_zh_abs_body - 8), -1):
            p = p_children[back]
            if not _ptext(p).strip():
                continue
            pPr_c = p.find(_qn_("w:pPr"))
            if pPr_c is None:
                break
            jc_c = pPr_c.find(_qn_("w:jc"))
            if jc_c is not None and jc_c.get(_qn_("w:val")) == "center":
                idx_zh_title = back
            break
    if idx_en_abs_body is not None:
        for back in range(idx_en_abs_body - 1, max(-1, idx_en_abs_body - 8), -1):
            p = p_children[back]
            if not _ptext(p).strip():
                continue
            pPr_c = p.find(_qn_("w:pPr"))
            if pPr_c is None:
                break
            jc_c = pPr_c.find(_qn_("w:jc"))
            if jc_c is not None and jc_c.get(_qn_("w:val")) == "center":
                idx_en_title = back
            break

    # 拼块区间
    if idx_toc_title is not None:
        end = idx_zh_title if idx_zh_title is not None else idx_h1_first or n
        blocks["toc"] = (idx_toc_title, end)
    if idx_zh_title is not None:
        end = idx_en_title if idx_en_title is not None else idx_h1_first or n
        blocks["zh_abstract"] = (idx_zh_title, end)
    if idx_en_title is not None:
        end = idx_h1_first or n
        blocks["en_abstract"] = (idx_en_title, end)
    if idx_h1_first is not None:
        end = idx_biblio_h1 if idx_biblio_h1 is not None else n
        blocks["body"] = (idx_h1_first, end)
    if idx_biblio_h1 is not None:
        end = idx_ack_h1 if idx_ack_h1 is not None else n
        blocks["bibliography"] = (idx_biblio_h1, end)
    if idx_ack_h1 is not None:
        end = idx_achieve_h1 if idx_achieve_h1 is not None else n
        blocks["acknowledge"] = (idx_ack_h1, end)
    if idx_achieve_h1 is not None:
        blocks["achievement"] = (idx_achieve_h1, n)
    return blocks


# ---------- §4 样式级规则 ----------


def _rule_styles(styles_el):
    """FORMAT_RULES.md §4 样式级规则。"""
    # Heading3: sz=24 (12pt 小四)
    for s in styles_el.findall(_qn_("w:style")):
        sid = s.get(_qn_("w:styleId"))
        if sid not in ("Heading3", "Heading3Char"):
            continue
        rPr = s.find(_qn_("w:rPr"))
        if rPr is None:
            rPr = _el_("w:rPr")
            s.append(rPr)
        for tag in ("w:sz", "w:szCs"):
            for old in list(rPr.findall(_qn_(tag))):
                rPr.remove(old)
            el = _el_(tag)
            el.set(_qn_("w:val"), "24")
            rPr.append(el)
    # TOC1/TOC2/TOC3: spacing after=0 line=360 auto + tabs right dot
    # 顶层设计（学校检测器规范）：
    #   TOC1 (一级章节项 "1 绪论"/"参考文献"/...)  → firstLine=0   左对齐，中文黑体
    #   TOC2 (二级章节项 "1.1 ..."/"1.2 ..."/...) → firstLine=480 缩进 2 字符（检测器强制）
    #   TOC3 (三级章节项 "1.1.1 ..."/...)         → firstLine=480 缩进 2 字符
    # 关键：用 firstLine 而非 left 表达缩进，检测器按段级 firstLine 判定。
    _TOC_FL_RULES = {
        "TOC1": ("0", "0"),
        "TOC2": ("200", "480"),
        "TOC3": ("200", "480"),
    }
    for s in styles_el.findall(_qn_("w:style")):
        sid = s.get(_qn_("w:styleId"))
        if sid not in _TOC_FL_RULES:
            continue
        pPr = s.find(_qn_("w:pPr"))
        if pPr is None:
            pPr = _el_("w:pPr")
            s.append(pPr)
        _set_spacing(pPr, before=0, after=0, line=360, lineRule="auto")
        ind = _ensure_child(pPr, "w:ind")
        # 清所有缩进属性，从干净状态开始
        for attr in ("leftChars", "left", "firstLineChars", "firstLine",
                     "rightChars", "right", "hanging", "hangingChars"):
            k = _qn_(f"w:{attr}")
            if k in ind.attrib:
                del ind.attrib[k]
        flc, fl = _TOC_FL_RULES[sid]
        ind.set(_qn_("w:firstLineChars"), flc)
        ind.set(_qn_("w:firstLine"), fl)
        # tab right dot
        tabs = _ensure_child(pPr, "w:tabs")
        has_right_dot = False
        for tab in tabs.findall(_qn_("w:tab")):
            if tab.get(_qn_("w:val")) == "right":
                tab.set(_qn_("w:leader"), "dot")
                if not tab.get(_qn_("w:pos")):
                    tab.set(_qn_("w:pos"), "8306")
                has_right_dot = True
        if not has_right_dot:
            tab_el = _el_("w:tab")
            tab_el.set(_qn_("w:val"), "right")
            tab_el.set(_qn_("w:leader"), "dot")
            tab_el.set(_qn_("w:pos"), "8306")
            tabs.append(tab_el)
    # Bibliography: spacing line=400 lineRule=exact
    for s in styles_el.findall(_qn_("w:style")):
        sid = s.get(_qn_("w:styleId"))
        if sid != "Bibliography":
            continue
        pPr = s.find(_qn_("w:pPr"))
        if pPr is None:
            pPr = _el_("w:pPr")
            s.append(pPr)
        _set_spacing(pPr, line=400, lineRule="exact")


# ---------- §2.1 toc 块 ----------


def _rule_toc(p_children, block):
    """FORMAT_RULES.md §2.1。
    - 目录标题"目  录"段：加粗 run
    - TOC1/TOC2 段的样式层属性由 _rule_styles 负责；段级不另设
    """
    if block is None:
        return
    start, end = block
    # 目录标题段（首段）
    p_title = p_children[start]
    for r_el in p_title.findall(_qn_("w:r")):
        _run_set_bold(r_el)


# ---------- §2.2 abstract 块 ----------


def _rule_abstract_page(p_children, block, lang):
    """FORMAT_RULES.md §2.2。论文题目段 spacing before=120 after=31 + 加粗。"""
    if block is None:
        return
    start, end = block
    # 第一段是论文题目段（居中 16pt）
    p_title = p_children[start]
    pPr = _ensure_pPr(p_title)
    _set_spacing(pPr, before=120, after=31, line=360, lineRule="auto")
    for r_el in p_title.findall(_qn_("w:r")):
        _run_set_bold(r_el)


# ---------- §2.3 body 块 ----------


def _rule_body(p_children, block):
    """FORMAT_RULES.md §2.3 正文。

    按段类型分派：H1/H2/H3 不动（样式控制），正文/caption/公式/代码 caption 分别处理。
    """
    if block is None:
        return
    start, end = block
    for i in range(start, end):
        p = p_children[i]
        style = _pstyle(p)
        if style.startswith("Heading"):
            continue
        txt = _ptext(p).strip()
        if not txt:
            continue
        pPr = _ensure_pPr(p)
        # 公式编号纯段 → 右对齐
        if _EQN_NUMBER_RE.match(txt):
            _set_jc(pPr, "right")
            continue
        # Image/Table Caption 样式段：样式已经居中无缩进，跳过（检测器对 caption 无首行缩进无异议）
        if style in ("ImageCaption", "TableCaption", "Caption"):
            _set_indent(pPr, firstLineChars=0, firstLine=0)
            continue
        # 代码/算法 caption 判定（根因级）：
        #   pandoc 把 \noindent\textbf{代码X-Y ...} 输出为整段所有 run 加粗。
        #   正文段 "算法5-1 给出..." 虽以"算法N-M"开头但 run 不加粗。
        #   所以真正的区分条件是"段内所有非空 run 都加粗" — 无需长度阈值。
        if (_CODE_ALGO_SHORT_CAP_RE.match(txt)
                and "的伪代码" not in txt
                and _all_runs_bold(p)):
            _set_indent(pPr, firstLineChars=0, firstLine=0)
            _set_jc(pPr, "center")
            continue
        # 其他以"图X.X/表X.X/代码X-Y/算法X-Y"开头的正文段 → 首行缩进 2 字符
        if _CAP_ANY_RE.match(txt):
            _set_indent(pPr, firstLineChars=200, firstLine=480)
            # "算法X-Y 的伪代码..." 等属于正文描述段：两端对齐（因 caption 处理流水可能误刷成居中）
            if "的伪代码" in txt:
                _set_jc(pPr, "both")
            continue


# ---------- §2.4 bibliography 块 ----------


def _rule_bibliography(p_children, block):
    """FORMAT_RULES.md §2.4。行距固定 20 磅 exact。"""
    if block is None:
        return
    start, end = block
    for i in range(start, end):
        p = p_children[i]
        txt = _ptext(p)
        style = _pstyle(p)
        if style == "Heading1":
            continue  # "参考文献" 标题由 Heading1 样式控制
        # Bibliography 样式段 + [N] 开头段
        if not (style.lower().startswith(("bibliograph", "reference")) or _BRACKET_REF_RE.match(txt)):
            continue
        pPr = _ensure_pPr(p)
        _set_spacing(pPr, line=400, lineRule="exact")


# ---------- §2.5 acknowledge 块 ----------


def _rule_acknowledge(p_children, block):
    """FORMAT_RULES.md §2.5。H1 样式控制 "致 谢" 标题，内容段 1.5 倍行距已由主处理链保证。"""
    # 无须额外处理
    return


# ---------- §2.6 achievement 块 ----------


def _rule_achievement(p_children, block):
    """FORMAT_RULES.md §2.6。成果页标题居中 + 正文按正文规则。"""
    if block is None:
        return
    start, end = block
    p_title = p_children[start]
    style = _pstyle(p_title)
    if style != "Heading1":
        pPr = _ensure_pPr(p_title)
        _set_jc(pPr, "center")


# ---------- §3 跨块字符级规则 ----------


def _strip_caption_prefix_space(p_el):
    """caption 段"图 N.M" / "表 N.M" / "代码 N-M" → "图N.M"（去所有前置空白）。"""
    ts = p_el.findall(".//" + _qn_("w:t"))
    if not ts:
        return False
    full = "".join((t.text or "") for t in ts)
    m = _CAP_LEAD_FIG_RE.match(full) or _CAP_LEAD_TAB_RE.match(full) or _CAP_LEAD_CODE_RE.match(full)
    if not m:
        return False
    del_start, del_end = m.start(2), m.end(2)
    if del_start >= del_end:
        return False
    pos = 0
    for t in ts:
        tt = t.text or ""
        if not tt:
            continue
        t_len = len(tt)
        t_end = pos + t_len
        lo = max(pos, del_start)
        hi = min(t_end, del_end)
        if lo < hi:
            rel_lo = lo - pos
            rel_hi = hi - pos
            t.text = tt[:rel_lo] + tt[rel_hi:]
        pos = t_end
    return True


def _rule_caption_prefix_space(body):
    """FORMAT_RULES.md §3.9 + §5.1-5.2：图/表/代码/算法 caption 前缀与数字之间无空格。"""
    for child in list(body):
        if child.tag != _qn_("w:p"):
            continue
        style_val = _pstyle(child)
        txt = _ptext(child).strip()
        if not txt:
            continue
        is_cap_style = style_val in ("ImageCaption", "TableCaption", "Caption")
        is_cap_by_text = _CAP_ANY_RE.match(txt) is not None and len(txt) < 120
        if is_cap_style or is_cap_by_text:
            _strip_caption_prefix_space(child)


def _rule_cjk_half_comma(doc):
    """FORMAT_RULES.md §3.5：中文段里半角 "," → 全角"，"。参考文献段跳过。"""
    for p in doc.paragraphs:
        style_name = p.style.name if p.style else ""
        if style_name.lower().startswith(("bibliograph", "reference")):
            continue
        if style_name in ("SourceCode", "Verbatim"):
            continue
        for r in p.runs:
            if not r.text:
                continue
            new_text = _CJK_COMMA_RE.sub(r"\1，\2", r.text)
            if new_text != r.text:
                r.text = new_text


def _rule_body_runlevel_bold(p_children, blocks):
    """FORMAT_RULES.md §3.6：正文段 run-level bold 清除。

    保留加粗的段：
      - Image/Table Caption 样式段（caption）
      - Heading 系列段
      - 代码/算法 caption 段（源自 \\noindent\\textbf{代码X-Y ...}，整段所有 run 加粗）
      - 摘要/Abstract/关键词/Key words 前缀段
      - 目录标题段、论文题目段（居中 16pt）
      - 参考文献标题段（H1）

    作用范围：body 块 + acknowledge 块 + achievement 块（需要清除 run 级 bold）
    """
    ranges = []
    for k in ("body", "acknowledge", "achievement", "bibliography"):
        if k in blocks:
            ranges.append(blocks[k])
    if not ranges:
        return
    for start, end in ranges:
        for i in range(start, end):
            p = p_children[i]
            style = _pstyle(p)
            if style.startswith("Heading"):
                continue
            if style in ("ImageCaption", "TableCaption", "Caption"):
                continue
            txt = _ptext(p).strip()
            if not txt:
                continue
            if txt.lstrip().startswith(("摘要", "关键词", "Abstract", "Key words", "Keywords")):
                continue
            # 代码/算法 caption 保留加粗（判定条件与 _rule_body 完全一致）
            if (_CODE_ALGO_SHORT_CAP_RE.match(txt)
                    and "的伪代码" not in txt
                    and _all_runs_bold(p)):
                continue
            # 纯参考文献条目（[N] 开头）不清 bold（通常无 bold 无所谓，清了也无害）
            for r_el in p.findall(_qn_("w:r")):
                _run_clear_bold(r_el)


def _rule_abstract_runlevel_bold(p_children, blocks):
    """摘要/Abstract 块内：论文题目段保留加粗；前缀段保留加粗；内容段清除 run-level bold。"""
    for k in ("zh_abstract", "en_abstract"):
        if k not in blocks:
            continue
        start, end = blocks[k]
        # 跳过第一段（论文题目段）
        for i in range(start + 1, end):
            p = p_children[i]
            txt = _ptext(p).strip()
            if not txt:
                continue
            # 摘要/关键词前缀段本身需要加粗，但只保留前缀 run 的 bold 已由 bolden_abstract_prefixes 处理
            # 这里不额外清除以免破坏前缀
            if txt.lstrip().startswith(("摘要", "关键词", "Abstract", "Key words", "Keywords")):
                continue
            for r_el in p.findall(_qn_("w:r")):
                _run_clear_bold(r_el)


# ---------- 主入口 ----------


def _apply_format_rules(doc):
    """FORMAT_RULES.md 规则驱动的统一格式处理器入口。"""
    from docx.oxml.ns import qn as _qn
    body = doc.element.body
    styles_el = doc.styles.element

    # 1. 样式级规则（Heading3 / TOC1/2 / Bibliography）
    _rule_styles(styles_el)

    # 2. 识别章节块
    p_children = [c for c in list(body) if c.tag == _qn("w:p")]
    blocks = _identify_blocks(p_children)

    # 3. 按块应用规则
    _rule_toc(p_children, blocks.get("toc"))
    _rule_abstract_page(p_children, blocks.get("zh_abstract"), "zh")
    _rule_abstract_page(p_children, blocks.get("en_abstract"), "en")
    _rule_body(p_children, blocks.get("body"))
    _rule_bibliography(p_children, blocks.get("bibliography"))
    _rule_acknowledge(p_children, blocks.get("acknowledge"))
    _rule_achievement(p_children, blocks.get("achievement"))

    # 4. 跨块字符级规则
    _rule_caption_prefix_space(body)
    _rule_cjk_half_comma(doc)
    _rule_body_runlevel_bold(p_children, blocks)
    _rule_abstract_runlevel_bold(p_children, blocks)


# 向后兼容：旧入口名继续暴露，实际跑新规则链
def _close_inspector_issues(doc) -> None:
    _apply_format_rules(doc)
