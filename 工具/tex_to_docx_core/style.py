from __future__ import annotations

import re

from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from .docx_common import _ensure_pPr, _ensure_child, _set_rfonts, _set_spacing, _set_indent


def bolden_abstract_prefixes(doc) -> None:
    """给摘要/Abstract 段首 "摘要：" "关键词：" "Abstract:" "Key words:" 前缀加粗。"""
    prefixes = [
        (re.compile(r"^(摘\s*要[：:])"), "黑体"),
        (re.compile(r"^(关键词[：:])"), "黑体"),
        (re.compile(r"^(Abstract\s*[:：])", re.IGNORECASE), "黑体"),
        (re.compile(r"^(Key\s*words?\s*[:：])", re.IGNORECASE), "黑体"),
    ]
    for p in doc.paragraphs:
        txt = (p.text or "").lstrip()
        for pat, cjk in prefixes:
            m = pat.match(txt)
            if not m:
                continue
            prefix = m.group(1)
            # 找到首个 run，拆分开头
            if not p.runs:
                break
            r0 = p.runs[0]
            if not r0.text:
                break
            # 如果首 run 就是前缀或包含前缀，直接给首 run 加粗并改字体
            if r0.text.startswith(prefix):
                # 插入新 run：bold 前缀；保留剩余普通 run
                remainder = r0.text[len(prefix) :]
                r0.text = prefix
                r0.bold = True
                if cjk:
                    _set_rfonts(r0, ascii_=r0.font.name or "Times New Roman", cjk=cjk)
                if remainder:
                    # 在 r0 后面插一个新 run 承载剩余
                    new_r = p.add_run(remainder)
                    # 把 new_r 的 XML 挪到 r0 之后
                    r0._element.addnext(new_r._element)
            break


def normalize_all_fonts(doc) -> None:
    """全文档字体正则化：
    1. styles.xml docDefaults 改为 中文宋体 / 英文 Times New Roman（摒弃 minorHAnsi 主题 → 大陆 Office 解析为等线）
    2. Normal 样式显式固定 rFonts
    3. 遍历所有 run，把 eastAsia/ascii 为 等线/Dengxian/Calibri/仿宋 / <none> 的强制改成 宋体 / TNR
       （标题和特殊段主动设的"黑体"保留）"""
    CJK_SAFE = {"宋体", "黑体", "仿宋", "楷体", "Times New Roman"}
    ASC_SAFE = {"Times New Roman"}

    # --- 1. 修 styles.xml ---
    styles_el = doc.styles.element  # <w:styles>

    # docDefaults/rPrDefault/rPr/rFonts
    docDef = styles_el.find(qn("w:docDefaults"))
    if docDef is not None:
        rPrDef = docDef.find(qn("w:rPrDefault"))
        if rPrDef is not None:
            rPr = rPrDef.find(qn("w:rPr"))
            if rPr is None:
                rPr = OxmlElement("w:rPr")
                rPrDef.append(rPr)
            old = rPr.find(qn("w:rFonts"))
            if old is not None:
                rPr.remove(old)
            rF = OxmlElement("w:rFonts")
            rF.set(qn("w:ascii"), "Times New Roman")
            rF.set(qn("w:hAnsi"), "Times New Roman")
            rF.set(qn("w:eastAsia"), "宋体")
            rF.set(qn("w:cs"), "Times New Roman")
            # 插入到 rPr 首位（rFonts 必须先于其他 rPr 子元素）
            rPr.insert(0, rF)

    # Normal 样式也显式写上 rFonts，防止某些 Word 忽略 docDefaults
    for style in styles_el.findall(qn("w:style")):
        sid = style.get(qn("w:styleId"))
        if sid == "Normal":
            rPr = style.find(qn("w:rPr"))
            if rPr is None:
                rPr = OxmlElement("w:rPr")
                style.append(rPr)
            old = rPr.find(qn("w:rFonts"))
            if old is not None:
                rPr.remove(old)
            rF = OxmlElement("w:rFonts")
            rF.set(qn("w:ascii"), "Times New Roman")
            rF.set(qn("w:hAnsi"), "Times New Roman")
            rF.set(qn("w:eastAsia"), "宋体")
            rF.set(qn("w:cs"), "Times New Roman")
            rPr.insert(0, rF)
        elif sid in ("Caption", "ImageCaption", "TableCaption"):
            # 根除 Caption 家族样式的 italic（pandoc 默认 Caption = italic，规范要求不斜）
            rPr = style.find(qn("w:rPr"))
            if rPr is not None:
                for i_tag in (qn("w:i"), qn("w:iCs")):
                    old_i = rPr.find(i_tag)
                    if old_i is not None:
                        rPr.remove(old_i)
                # 显式关闭 italic（子样式就算叠加也不会回来）
                neg_i = OxmlElement("w:i")
                neg_i.set(qn("w:val"), "0")
                rPr.append(neg_i)

    # --- 2. 修 theme1.xml（兜底）---
    try:
        theme_part = None
        for rel in doc.part.rels.values():
            if rel.reltype.endswith("/theme"):
                theme_part = rel.target_part
                break
        if theme_part is not None:
            import re as _re

            xml = theme_part.blob.decode("utf-8")

            # minorFont 里 typeface="Calibri" / latin typeface 改成 TNR；东亚 typeface 改成宋体
            # 粗略替换：在 <a:minorFont> 段内的 typeface
            def _sub_minor(m):
                block = m.group(0)
                block = _re.sub(
                    r'<a:latin typeface="[^"]*"',
                    '<a:latin typeface="Times New Roman"',
                    block,
                )
                block = _re.sub(
                    r'<a:ea typeface="[^"]*"', '<a:ea typeface="宋体"', block
                )
                # 东亚语言 script 的 typeface 替换
                block = _re.sub(
                    r'(<a:font script="Hans" typeface=")[^"]*(")', r"\1宋体\2", block
                )
                block = _re.sub(
                    r'(<a:font script="Hant" typeface=")[^"]*(")', r"\1宋体\2", block
                )
                return block

            xml2 = _re.sub(
                r"<a:minorFont>.*?</a:minorFont>", _sub_minor, xml, flags=_re.DOTALL
            )
            if xml2 != xml:
                theme_part._blob = xml2.encode("utf-8")
    except Exception:
        pass  # theme 修改失败不影响主流程

    # --- 3. 遍历所有 run，强制覆盖字体 ---
    def _is_inside_code_table(r_el) -> bool:
        parent = r_el.getparent()
        while parent is not None:
            if parent.tag == qn("w:tbl"):
                all_t = "".join(t.text or "" for t in parent.findall(".//" + qn("w:t")))
                return any(
                    marker in all_t
                    for marker in ("//", "/*", "#include", "{", "}", ";", "==", "->")
                )
            parent = parent.getparent()
        return False

    def _fix_run(r_el):
        if _is_inside_code_table(r_el):
            return
        rPr = r_el.find(qn("w:rPr"))
        if rPr is None:
            rPr = OxmlElement("w:rPr")
            r_el.insert(0, rPr)
        rF = rPr.find(qn("w:rFonts"))
        if rF is None:
            rF = OxmlElement("w:rFonts")
            rPr.insert(0, rF)
        # 清掉 themeFont 属性（它们会被 theme 字体解析）
        for attr in (
            qn("w:asciiTheme"),
            qn("w:hAnsiTheme"),
            qn("w:eastAsiaTheme"),
            qn("w:cstheme"),
        ):
            if attr in rF.attrib:
                del rF.attrib[attr]
        # eastAsia：如果不是宋体/黑体/仿宋/楷体/TNR，就改成宋体
        cjk = rF.get(qn("w:eastAsia"))
        if cjk not in CJK_SAFE:
            rF.set(qn("w:eastAsia"), "宋体")
        # ascii / hAnsi：如果不是 TNR 就改成 TNR
        for attr_name in ("ascii", "hAnsi", "cs"):
            v = rF.get(qn(f"w:{attr_name}"))
            if v not in ASC_SAFE:
                rF.set(qn(f"w:{attr_name}"), "Times New Roman")

    body = doc.element.body
    for r_el in body.iter(qn("w:r")):
        _fix_run(r_el)

    # footer 也过一遍
    for sect in doc.sections:
        for p in sect.footer.paragraphs:
            for r_el in p._element.iter(qn("w:r")):
                _fix_run(r_el)


def enable_latin_word_break(doc) -> None:
    """允许 Latin 单词在行末从中间断开（东亚排版场景）。
    w:wordWrap val="0" 语义：关闭"在单词边界换行"的限制，允许 PathBoundsDecider
    这类无连字符驼峰词在任意字符处换行。双保险：
      1) Normal 样式 pPr（所有段落继承）
      2) settings.xml doNotWrapTextWithPunct 家族兜底"""
    styles_el = doc.styles.element
    for style in styles_el.findall(qn("w:style")):
        if style.get(qn("w:styleId")) != "Normal":
            continue
        pPr = style.find(qn("w:pPr"))
        if pPr is None:
            pPr = OxmlElement("w:pPr")
            # pPr 必须在 rPr 之前（w:style 子元素顺序）
            rPr = style.find(qn("w:rPr"))
            if rPr is not None:
                rPr.addprevious(pPr)
            else:
                style.append(pPr)
        old = pPr.find(qn("w:wordWrap"))
        if old is not None:
            pPr.remove(old)
        ww = OxmlElement("w:wordWrap")
        ww.set(qn("w:val"), "0")
        pPr.append(ww)
        break

    # 文档级 compat 兜底：部分 Word 版本读 settings.xml 的 compat 开关
    settings_el = doc.settings.element
    compat = settings_el.find(qn("w:compat"))
    if compat is None:
        compat = OxmlElement("w:compat")
        settings_el.append(compat)
    # doNotUseEastAsianBreakRules 会让 Word 用西文断行规则（含 word boundary）
    # 存在时必须去掉，否则 wordWrap=0 被忽略
    bad = compat.find(qn("w:doNotUseEastAsianBreakRules"))
    if bad is not None:
        compat.remove(bad)


def normalize_paragraph_spacing(doc) -> None:
    """正文段规格全局拉通：
    - 所有段（含表内）：snapToGrid=false，避免 Word "对齐到网格" 压缩行间距
    - 正文段（非标题/非表内/非代码块/非签名页）：
        * 段前 6 磅 (before=120 twips)
        * 段后 0 磅 (after=0 twips)
        * 首行缩进 2 字符 (firstLineChars=200)，去掉 firstLine 绝对 twips（避免字符单位被覆盖）
    - 标题段：段前/段后 0 磅，snapToGrid=false
    - 摘要中英文题目页标题：段前 6 磅，段后 1.54 磅
    - 表格内段、代码块段、签名 tab 段：只设 snapToGrid=false，不动其他"""
    body = doc.element.body
    # 封面区结束位置：第一个 sectPr
    cover_end = 0
    for i, child in enumerate(body):
        if child.tag == qn("w:p"):
            pPr = child.find(qn("w:pPr"))
            if pPr is not None and pPr.find(qn("w:sectPr")) is not None:
                cover_end = i
                break
        elif child.tag == qn("w:sectPr"):
            cover_end = i
            break

    heading_prefixes = ("Heading", "Title", "TOC")

    def _set_snap_false(pPr):
        snap = pPr.find(qn("w:snapToGrid"))
        if snap is None:
            snap = OxmlElement("w:snapToGrid")
            pPr.append(snap)
        snap.set(qn("w:val"), "0")

    def _set_spacing_local(pPr, before=None, after=None):
        """局部快捷：只处理 before/after；line/lineRule 由 docx_common._set_spacing 负责。"""
        spacing = pPr.find(qn("w:spacing"))
        if spacing is None:
            spacing = OxmlElement("w:spacing")
            pPr.append(spacing)
        if before is not None:
            spacing.set(qn("w:before"), str(before))
            spacing.set(qn("w:beforeLines"), "0")
            spacing.set(qn("w:beforeAutospacing"), "0")
        if after is not None:
            spacing.set(qn("w:after"), str(after))
            spacing.set(qn("w:afterLines"), "0")
            spacing.set(qn("w:afterAutospacing"), "0")

    def _set_firstline_chars(pPr):
        ind = pPr.find(qn("w:ind"))
        if ind is None:
            ind = OxmlElement("w:ind")
            pPr.append(ind)
        ind.set(qn("w:firstLineChars"), "200")
        # 去掉 firstLine 绝对值，优先 chars；Word 会按 chars * 当前字号重算
        for attr in ("firstLine",):
            k = qn(f"w:{attr}")
            if k in ind.attrib:
                del ind.attrib[k]

    # 1. 处理顶层段
    for i, child in enumerate(body):
        if child.tag != qn("w:p"):
            continue
        pPr = _ensure_pPr(child)
        _set_snap_false(pPr)

        pStyle = pPr.find(qn("w:pStyle"))
        style_val = pStyle.get(qn("w:val")) if pStyle is not None else ""
        is_heading = any(style_val.startswith(pre) for pre in heading_prefixes)
        txt = "".join((t.text or "") for t in child.findall(".//" + qn("w:t")))

        if is_heading:
            # 标题段：报告要求标题段段前/段后归零
            _set_spacing(pPr, before=0, after=0)
            continue

        # 参考文献条目（pandoc citeproc 样式 Bibliography / References）：
        # 不加首行缩进；清掉可能残留的 firstLine / firstLineChars
        if style_val.lower().startswith(("bibliograph", "reference")):
            ind = pPr.find(qn("w:ind"))
            if ind is not None:
                for attr in ("firstLine", "firstLineChars"):
                    k = qn(f"w:{attr}")
                    if k in ind.attrib:
                        del ind.attrib[k]
            continue

        # 跳过封面区、含 sectPr 的段
        if i <= cover_end:
            continue

        # 跳过代码块表格前的 caption 段（保持居中样式）
        jc = pPr.find(qn("w:jc"))
        jc_val = jc.get(qn("w:val")) if jc is not None else None
        # 居中/右对齐段：图题/表题/公式 —— 不加首行缩进，但也设 6/1.54 间距
        if jc_val in ("center", "right", "distribute"):
            # "目  录"标题段特殊：段前 1.5 行，段后 0.5 行
            txt_c = "".join(
                (t.text or "") for t in child.findall(".//" + qn("w:t"))
            ).strip()
            if txt_c in ("目  录", "目录"):
                spacing = pPr.find(qn("w:spacing"))
                if spacing is None:
                    spacing = OxmlElement("w:spacing")
                    pPr.append(spacing)
                spacing.set(qn("w:beforeLines"), "150")
                spacing.set(qn("w:afterLines"), "50")
                for attr in (
                    "before",
                    "after",
                    "beforeAutospacing",
                    "afterAutospacing",
                ):
                    k = qn(f"w:{attr}")
                    if k in spacing.attrib:
                        del spacing.attrib[k]
            elif _looks_like_abstract_title(txt_c):
                _set_spacing(pPr, before=0, after=0)
            elif style_val in ("ImageCaption",) or (
                style_val == "Caption" and txt_c.startswith("图")
            ):
                # 图题：段前段后 0（学校模板要求 caption 四围归零）
                _set_spacing(pPr, before=0, after=0)
            elif style_val in ("TableCaption",) or (
                style_val == "Caption" and txt_c.startswith("表")
            ):
                # 表题：段前段后 0（学校模板要求 caption 四围归零）
                _set_spacing(pPr, before=0, after=0)
            elif style_val == "CaptionedFigure":
                # pandoc 把「图+图题」整体包成 CaptionedFigure 段；学校模板
                # 要求段前段后 0（跨页留白由 1.5 倍行距与浮动体自身承载）
                _set_spacing(pPr, before=0, after=0)
            else:
                _set_spacing(pPr, before=120, after=0)
            continue

        # 跳过签名行（tabs 已定义，firstLineChars 会破坏对齐）
        tabs = pPr.find(qn("w:tabs"))
        if tabs is not None and ("签名" in txt or "日期" in txt):
            # 清掉首行缩进（签名行不需要）
            ind = pPr.find(qn("w:ind"))
            if ind is not None:
                pPr.remove(ind)
            continue

        # 摘要/Abstract 首段（含正文）：学校模板要求段前段后 0
        # 触发词：以"摘要："或"Abstract"开头（兼容半角/全角冒号）
        txt_strip = txt.lstrip()
        if txt_strip.startswith(("摘要：", "摘要:", "Abstract:", "Abstract：")) or (
            txt_strip.startswith("Abstract") and len(txt_strip) > 8 and txt_strip[8] in (" ", ":", "：")
        ):
            _set_spacing(pPr, before=0, after=0)
            _set_firstline_chars(pPr)
            continue

        # 常规正文段：段前 0、段后 0、首行缩进 2 字符
        # （正文间距由固定 1.5 倍行距承载，不再额外加段前/段后）
        _set_spacing(pPr, before=0, after=0)
        _set_firstline_chars(pPr)

    # 2. 处理表格内所有段（只设 snapToGrid=false，不动 spacing/ind）
    for tbl in doc.element.body.iter(qn("w:tbl")):
        for p_el in tbl.iter(qn("w:p")):
            pPr = _ensure_pPr(p_el)
            _set_snap_false(pPr)


def justify_body_paragraphs(doc) -> None:
    """正文段两端对齐（jc=both）。跳过：表格内段、居中/右对齐段、标题段、封面/签名页特殊段。"""
    body = doc.element.body
    # 封面区结束位置：第一个 sectPr 之前的段都视为封面
    cover_end = 0
    for i, child in enumerate(body):
        if child.tag == qn("w:p"):
            pPr = child.find(qn("w:pPr"))
            if pPr is not None and pPr.find(qn("w:sectPr")) is not None:
                cover_end = i
                break
        elif child.tag == qn("w:sectPr"):
            cover_end = i
            break

    # 标题样式关键字
    heading_style_prefixes = ("Heading", "Title", "TOC")

    for i, child in enumerate(body):
        if child.tag != qn("w:p"):
            continue
        # 表格内段？iter 顶层 body，这里 child 都是顶层段，不会在表格内
        pPr = child.find(qn("w:pPr"))
        if pPr is None:
            pPr = OxmlElement("w:pPr")
            child.insert(0, pPr)
        # 样式名
        pStyle = pPr.find(qn("w:pStyle"))
        style_val = pStyle.get(qn("w:val")) if pStyle is not None else ""
        if any(style_val.startswith(pre) for pre in heading_style_prefixes):
            continue
        # 已有 center / right / distribute：跳过
        jc = pPr.find(qn("w:jc"))
        if jc is not None and jc.get(qn("w:val")) in ("center", "right", "distribute"):
            continue
        # 跳过封面区
        if i <= cover_end:
            continue
        # 跳过含分页符的空段（首行有 br type=page）
        brs = child.findall(".//" + qn("w:br"))
        if (
            brs
            and not "".join(
                (t.text or "") for t in child.findall(".//" + qn("w:t"))
            ).strip()
        ):
            # 纯分页符段
            pass  # 还是设 jc=both 不影响
        # 设 both
        if jc is None:
            jc = OxmlElement("w:jc")
            pPr.append(jc)
        jc.set(qn("w:val"), "both")


def fold_abstract_heading_into_body(doc) -> None:
    """删除 H1 "摘要" / "Abstract" 独立标题段 —— 它们只在管道里做结构锚点用
    （insert_thesis_title_pages / add_page_breaks_before_h1 / setup_page_numbers_and_sections
    定位都已完成）。删除后，摘要版面就是"题目（三号黑体加粗居中）+ 正文首段含【粗体 摘要：】前缀"。
    段首的 \\textbf{摘要：} / \\textbf{Abstract:} 已由 pandoc 转成粗体 run，
    bolden_abstract_prefixes 再把字体调成黑体加粗——不需要再手动插。
    约束：
    - 只删"纯文本 == '摘要' 或 'Abstract'"的 H1 段（避免误伤）
    - 保留上一段（可能是刚插的论文题目段）以及下一段（摘要正文首段）不动"""
    TARGETS = {"摘要", "Abstract"}
    to_remove = []
    for p in doc.paragraphs:
        if not (p.style and p.style.name == "Heading 1"):
            continue
        if (p.text or "").strip() not in TARGETS:
            continue
        to_remove.append(p._element)
    for el in to_remove:
        parent = el.getparent()
        if parent is not None:
            parent.remove(el)


def _looks_like_abstract_title(text: str) -> bool:
    """中英文摘要页插入的论文题目是居中段，但不属于目录/图题/公式。"""
    if not text:
        return False
    if text.startswith(("图", "表", "代码", "（")) or text in ("目  录", "目录"):
        return False
    if len(text) > 220:
        return False
    return bool(
        re.search(
            r"研究|设计|实现|系统|算法|框架|analysis|research|design|implementation|algorithm|planning",
            text,
            re.I,
        )
    )


def normalize_text_punctuation(doc) -> None:
    """修正常见正文标点误报：中文语境中的半角单双引号改成全角。

    只处理包含 CJK 字符的普通文本 run；代码块样式和字段指令不动，避免破坏代码清单与 TOC 域。
    """

    def _has_cjk(s: str) -> bool:
        return bool(re.search(r"[\u4e00-\u9fff]", s))

    def _replace_quotes(s: str) -> str:
        out: list[str] = []
        double_open = True
        single_open = True
        for ch in s:
            if ch == '"':
                out.append("“" if double_open else "”")
                double_open = not double_open
            elif ch == "'":
                out.append("‘" if single_open else "’")
                single_open = not single_open
            else:
                out.append(ch)
        return "".join(out)

    for p in doc.paragraphs:
        style_name = p.style.name if p.style else ""
        if style_name in {"SourceCode", "Code", "Verbatim"}:
            continue
        if not _has_cjk(p.text or ""):
            continue
        for run in p.runs:
            if run.text and ('"' in run.text or "'" in run.text):
                run.text = _replace_quotes(run.text)


def _ensure_toc_styles(doc) -> None:
    """把 TOC 1 设为黑体加粗、TOC 2 设为宋体不加粗，并给 TOC 2 缩进 2 字符。
    Word/WPS 更新目录域时，按这两个样式渲染一级/二级目录项。
    若模板没有对应样式则新建基于 Normal 的条目。"""
    styles_el = doc.styles.element

    def _find_or_create(style_id_candidates: tuple[str, ...], name_val: str):
        """按 styleId 候选集或 name 找现有样式；找不到则基于 Normal 新建。"""
        for style in styles_el.findall(qn("w:style")):
            sid = style.get(qn("w:styleId")) or ""
            if sid in style_id_candidates:
                return style
            name_el = style.find(qn("w:name"))
            if name_el is not None and (
                name_el.get(qn("w:val")) or ""
            ).lower() == name_val.lower():
                return style
        new = OxmlElement("w:style")
        new.set(qn("w:type"), "paragraph")
        new.set(qn("w:styleId"), style_id_candidates[0])
        name_el = OxmlElement("w:name")
        name_el.set(qn("w:val"), name_val)
        new.append(name_el)
        based = OxmlElement("w:basedOn")
        based.set(qn("w:val"), "Normal")
        new.append(based)
        next_el = OxmlElement("w:next")
        next_el.set(qn("w:val"), "Normal")
        new.append(next_el)
        styles_el.append(new)
        return new

    def _force_toc_style(target, *, bold: bool, cjk_font: str) -> None:
        """切断父链继承：显式写 bold 标志，顺带锁 cjk_font + 小四。
        ``cjk_font`` 用于 eastAsia 字体：TOC1 规范为黑体，TOC2 规范为宋体。"""
        rPr = target.find(qn("w:rPr"))
        if rPr is None:
            rPr = OxmlElement("w:rPr")
            target.append(rPr)
        # rFonts：cjk_font（中文）+ Times New Roman（西文）
        rFonts = rPr.find(qn("w:rFonts"))
        if rFonts is None:
            rFonts = OxmlElement("w:rFonts")
            rPr.insert(0, rFonts)
        rFonts.set(qn("w:eastAsia"), cjk_font)
        rFonts.set(qn("w:ascii"), "Times New Roman")
        rFonts.set(qn("w:hAnsi"), "Times New Roman")
        rFonts.set(qn("w:cs"), "Times New Roman")
        # 清除 theme-based font 属性，防止 theme 覆盖具体值
        for attr in ("w:asciiTheme", "w:hAnsiTheme", "w:eastAsiaTheme", "w:cstheme"):
            if rFonts.get(qn(attr)) is not None:
                del rFonts.attrib[qn(attr)]
        # 小四 (12pt = sz 24)
        for tag, val in (("w:sz", "24"), ("w:szCs", "24")):
            el = rPr.find(qn(tag))
            if el is None:
                el = OxmlElement(tag)
                rPr.append(el)
            el.set(qn("w:val"), val)
        # 显式设置加粗状态，同时关斜体，避免 basedOn 或模板残留影响
        for tag_name in ("w:b", "w:bCs", "w:i", "w:iCs"):
            for old in rPr.findall(qn(tag_name)):
                rPr.remove(old)
            neg = OxmlElement(tag_name)
            if tag_name in ("w:b", "w:bCs"):
                if not bold:
                    neg.set(qn("w:val"), "0")
            else:
                neg.set(qn("w:val"), "0")
            rPr.append(neg)
        # color 强制黑色，覆盖 basedOn 链里可能继承的 themeColor
        color = rPr.find(qn("w:color"))
        if color is None:
            color = OxmlElement("w:color")
            rPr.append(color)
        color.set(qn("w:val"), "000000")
        for attr in ("w:themeColor", "w:themeShade", "w:themeTint"):
            if color.get(qn(attr)) is not None:
                del color.attrib[qn(attr)]

    def _ensure_style_pPr(target):
        """专门给 <w:style> 元素用：pPr 必须放在 rPr 之前，否则 Word 拒绝读取样式。
        不同于 docx_common._ensure_pPr（那是段落用的 insert(0)）。
        """
        pPr = target.find(qn("w:pPr"))
        if pPr is None:
            pPr = OxmlElement("w:pPr")
            rPr = target.find(qn("w:rPr"))
            if rPr is not None:
                target.insert(list(target).index(rPr), pPr)
            else:
                target.append(pPr)
        return pPr

    def _ensure_leader_dot_tab(pPr):
        """给 pPr 写入带 dot leader 的右对齐 tab，页码对齐位置 8306 twips。
        用正文栏宽 A4-3-2=15.7cm ≈ 8906 twips，减去 600 微调 8306。"""
        tabs = pPr.find(qn("w:tabs"))
        if tabs is None:
            tabs = OxmlElement("w:tabs")
            pPr.insert(0, tabs)
        has_right_dot = False
        for tab in tabs.findall(qn("w:tab")):
            if tab.get(qn("w:val")) == "right":
                tab.set(qn("w:leader"), "dot")
                tab.set(qn("w:pos"), "8306")
                has_right_dot = True
        if not has_right_dot:
            tab_el = OxmlElement("w:tab")
            tab_el.set(qn("w:val"), "right")
            tab_el.set(qn("w:leader"), "dot")
            tab_el.set(qn("w:pos"), "8306")
            tabs.append(tab_el)

    # TOC 1（一级目录项）：黑体加粗；不设缩进；右侧 tab 带 dot leader
    toc1 = _find_or_create(("TOC1", "toc1"), "toc 1")
    _force_toc_style(toc1, bold=True, cjk_font="黑体")
    _ensure_leader_dot_tab(_ensure_style_pPr(toc1))

    # TOC 2（二级目录项）：宋体不加粗 + 首行缩进 2 字符（学校规范）
    toc2 = _find_or_create(("TOC2", "toc2"), "toc 2")
    _force_toc_style(toc2, bold=False, cjk_font="宋体")

    pPr = _ensure_style_pPr(toc2)
    _ensure_leader_dot_tab(pPr)
    ind = pPr.find(qn("w:ind"))
    if ind is None:
        ind = OxmlElement("w:ind")
        pPr.append(ind)
    # 用 firstLine 表达缩进，不用 left（与段级 ind 规则、_TOC_FL_RULES 严格对齐）
    for attr in ("leftChars", "left", "hanging", "hangingChars"):
        k = qn(f"w:{attr}")
        if k in ind.attrib:
            del ind.attrib[k]
    ind.set(qn("w:firstLineChars"), "200")
    ind.set(qn("w:firstLine"), "480")


def _ensure_toc2_not_bold(doc) -> None:
    """兼容旧名字的封装。"""
    _ensure_toc_styles(doc)


def normalize_toc_entries(doc) -> None:
    """目录项按检测报告口径清缩进、清段后；目录标题保留 1.5/0.5 行。

    Word/WPS 更新目录域后会生成 TOC1/TOC2/TOC3 样式段。当前脚本生成阶段只有
    目录域占位，不能用“目录标题之后到 1 绪论之前”这种范围判断，否则会误伤摘要页。
    """
    _ensure_toc_styles(doc)
    for p in doc.paragraphs:
        text = (p.text or "").strip()
        if text in ("目  录", "目录"):
            pPr = p._element.get_or_add_pPr()
            spacing = pPr.find(qn("w:spacing"))
            if spacing is None:
                spacing = OxmlElement("w:spacing")
                pPr.append(spacing)
            spacing.set(qn("w:beforeLines"), "150")
            spacing.set(qn("w:afterLines"), "50")
            for attr in ("before", "after", "beforeAutospacing", "afterAutospacing"):
                key = qn(f"w:{attr}")
                if key in spacing.attrib:
                    del spacing.attrib[key]
            continue

        style_name = p.style.name if p.style else ""
        pPr = p._element.get_or_add_pPr()
        pStyle = pPr.find(qn("w:pStyle"))
        style_val = pStyle.get(qn("w:val")) if pStyle is not None else ""
        if not (
            style_name.upper().startswith("TOC") or style_val.upper().startswith("TOC")
        ):
            continue
        spacing = pPr.find(qn("w:spacing"))
        if spacing is None:
            spacing = OxmlElement("w:spacing")
            pPr.append(spacing)
        spacing.set(qn("w:after"), "0")
        spacing.set(qn("w:afterLines"), "0")
        spacing.set(qn("w:afterAutospacing"), "0")
        ind = pPr.find(qn("w:ind"))
        if ind is None:
            ind = OxmlElement("w:ind")
            pPr.append(ind)
        for attr in ("left", "leftChars", "hanging", "hangingChars", "firstLine", "firstLineChars"):
            key = qn(f"w:{attr}")
            if key in ind.attrib:
                del ind.attrib[key]
        # 顶层设计（学校检测器规范）：
        #   TOC1 无首行缩进（左对齐）
        #   TOC2/TOC3 首行缩进 2 字符（firstLineChars=200, firstLine=480）
        # 段级覆盖样式级，必须与 postprocess._rule_styles 的 _TOC_FL_RULES 完全对齐
        style_upper = style_val.upper()
        style_lower = style_name.strip().lower()
        is_toc1 = (
            style_upper == "TOC1"
            or style_lower in ("toc 1", "toc1")
        )
        if is_toc1:
            ind.set(qn("w:firstLineChars"), "0")
            ind.set(qn("w:firstLine"), "0")
        else:
            ind.set(qn("w:firstLineChars"), "200")
            ind.set(qn("w:firstLine"), "480")


_CJK_RANGES = (
    "⺀-鿿"
    "豈-﫿"
    "︰-﹏"
    "\U00020000-\U0002fa1f"
)
_RE_CJK_SPACE_LATIN = re.compile(
    rf"([{_CJK_RANGES}])\s+([A-Za-z0-9Ͱ-Ͽ([\-])"
)
_RE_LATIN_SPACE_CJK = re.compile(
    rf"([A-Za-z0-9Ͱ-Ͽ)\]%.])\s+([{_CJK_RANGES}])"
)


_RE_CJK_TAIL = re.compile(rf"[{_CJK_RANGES}]$")
_RE_CJK_HEAD = re.compile(rf"^[{_CJK_RANGES}]")
_RE_LATIN_TAIL = re.compile(r"[A-Za-z0-9\u0370-\u03ff)\]%.]$")
_RE_LATIN_HEAD = re.compile(r"^[A-Za-z0-9\u0370-\u03ff([\-]")


def _strip_cjk_latin_in_para(p) -> None:
    """Strip CJK-Latin spaces in a single paragraph's runs."""
    runs = p.runs
    # Pass 1: within-run
    for run in runs:
        if not run.text:
            continue
        t = run.text
        t = _RE_CJK_SPACE_LATIN.sub(r"\1\2", t)
        t = _RE_LATIN_SPACE_CJK.sub(r"\1\2", t)
        if t != run.text:
            run.text = t
    # Pass 2: cross-run boundaries (trailing/leading space)
    for i in range(len(runs) - 1):
        cur = runs[i].text or ""
        nxt = runs[i + 1].text or ""
        if not cur or not nxt:
            continue
        if cur != cur.rstrip():
            stripped = cur.rstrip()
            if (
                (_RE_CJK_TAIL.search(stripped) and _RE_LATIN_HEAD.match(nxt))
                or (_RE_LATIN_TAIL.search(stripped) and _RE_CJK_HEAD.match(nxt))
            ):
                runs[i].text = stripped
        cur = runs[i].text or ""
        nxt = runs[i + 1].text or ""
        if nxt != nxt.lstrip():
            stripped_nxt = nxt.lstrip()
            if (
                (_RE_CJK_TAIL.search(cur) and _RE_LATIN_HEAD.match(stripped_nxt))
                or (_RE_LATIN_TAIL.search(cur) and _RE_CJK_HEAD.match(stripped_nxt))
            ):
                runs[i + 1].text = stripped_nxt
    # Pass 3: space-only run between CJK and Latin runs
    for i in range(len(runs) - 2):
        mid = runs[i + 1].text or ""
        if not mid or mid.strip():
            continue  # not a space-only run
        prev = runs[i].text or ""
        nxt = runs[i + 2].text or ""
        if not prev or not nxt:
            continue
        if (
            (_RE_CJK_TAIL.search(prev) and _RE_LATIN_HEAD.match(nxt))
            or (_RE_LATIN_TAIL.search(prev) and _RE_CJK_HEAD.match(nxt))
        ):
            runs[i + 1].text = ""


def strip_cjk_latin_spaces(doc) -> None:
    """删除中英文之间的手动空格，让 Word autoSpaceDE/DN 自动管理间距。"""
    _SKIP = ("toc", "bibliograph", "reference", "source", "verbatim", "heading",
             "image caption", "table caption")
    for p in doc.paragraphs:
        style_name = (p.style.name if p.style else "").lower()
        if style_name.startswith(_SKIP):
            continue
        _strip_cjk_latin_in_para(p)
    # Also process table cells (cover page, etc.)
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    style_name = (p.style.name if p.style else "").lower()
                    if style_name.startswith(_SKIP):
                        continue
                    _strip_cjk_latin_in_para(p)


def normalize_bibliography_text(doc) -> None:
    """修复参考文献英文半角小圆点后缺空格这类低风险格式问题。
    同时把 pandoc citeproc 在 "[N]" 后默认插入的 Tab（制表符）替换为单空格，
    避免编号与条目之间出现宽空白。"""
    for p in doc.paragraphs:
        style_name = p.style.name if p.style else ""
        if not style_name.lower().startswith(("bibliograph", "reference")):
            continue
        for run in p.runs:
            if run.text:
                # 1) 英文半角点后缺空格：修复为 ". X"
                run.text = re.sub(r"\.([A-Za-z])", r". \1", run.text)
                # 2) 去掉 Tab：替换为单空格
                if "\t" in run.text:
                    run.text = run.text.replace("\t", " ")
        # 3) 合并 "[N]" 编号 run 与其后的连续空白 run 成为 "[N] "，
        #    避免 "空格 + Tab" 两个空白 run 叠加产生大片宽白
        text = p.text
        m = re.match(r"\[\d+\]\s+", text)
        if not m:
            continue
        # 找到编号占据的 run 序列，重写为一个干净 "[N] " + 剩余文本
        # 直接操作底层 XML：把首个匹配 "[N]" 的 run 文本扩成 "[N] "，后续前置空白 run 清空
        runs = p.runs
        if not runs:
            continue
        first = runs[0]
        bracket_match = re.match(r"^(\[\d+\])\s*", first.text or "")
        if not bracket_match:
            continue
        first.text = bracket_match.group(1) + " "
        # 后续连续"纯空白"run 清空（避免累积）
        for r in runs[1:]:
            t = r.text or ""
            if t.strip() == "":
                r.text = ""
            else:
                # 若这个 run 文本以空白开头，去掉它开头的所有空白
                r.text = t.lstrip()
                break
