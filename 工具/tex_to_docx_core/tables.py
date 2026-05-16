from __future__ import annotations

import re

from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

from .config import FONT_SIZE
from .docx_common import _set_rfonts


def _is_code_block_table(tbl) -> bool:
    """代码块表格判定：pandoc 把 lstlisting verbatim 渲染成 1x1 单元格，里面每行一个段落或一个段内用 <w:br> 换行。
    特征：单行单列 且 (cell 内段数 >= 2 或 段内 br 数 >= 2) 且文本含代码符号。"""
    rows = tbl._element.findall(qn("w:tr"))
    if len(rows) != 1:
        return False
    tcs = rows[0].findall(qn("w:tc"))
    if len(tcs) != 1:
        return False
    all_t = "".join((t.text or "") for t in tcs[0].findall(".//" + qn("w:t")))
    ps = tcs[0].findall(qn("w:p"))
    brs = tcs[0].findall(".//" + qn("w:br"))
    # 段数 >= 2 或 br 数 >= 2
    if len(ps) < 2 and len(brs) < 2:
        return False
    # 代码特征
    code_markers = ("//", "/*", "#include", "{", "}", ";", "==", "->", "← ", "←")
    return any(m in all_t for m in code_markers)


def _clear_table_borders(tbl) -> None:
    """清掉表格所有边框（table 级 + cell 级），用于子图布局表/代码块表。"""
    t_el = tbl._element
    tblPr = t_el.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        t_el.insert(0, tblPr)
    # 表级边框全 nil
    old_tb = tblPr.find(qn("w:tblBorders"))
    if old_tb is not None:
        tblPr.remove(old_tb)
    tblBorders = OxmlElement("w:tblBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"), "nil")
        tblBorders.append(b)
    tblPr.append(tblBorders)
    # 每个 cell 边框也清
    for tr in t_el.findall(qn("w:tr")):
        for tc in tr.findall(qn("w:tc")):
            tcPr = tc.find(qn("w:tcPr"))
            if tcPr is None:
                continue
            old = tcPr.find(qn("w:tcBorders"))
            if old is not None:
                tcPr.remove(old)
            tcBorders = OxmlElement("w:tcBorders")
            for side in ("top", "left", "bottom", "right"):
                b = OxmlElement(f"w:{side}")
                b.set(qn("w:val"), "nil")
                tcBorders.append(b)
            tcPr.append(tcBorders)


def apply_three_line_tables(doc) -> None:
    """规范：表格一律三线表 + 整表居中。
    实现：
    1. tblPr 加 jc=center 整表居中；
    2. 清掉所有 tblBorders + tcBorders；
    3. 重设 tblBorders：top 粗 / bottom 粗 / left right insideH insideV 全 nil；
    4. 首行（header）每个 tc 底部加 single 4 细线 → 三线表第三线。"""

    # OOXML tblBorders/tcBorders 子元素顺序：top, left, bottom, right, insideH, insideV
    def _mkborder(side: str, val: str, sz: str | None = None):
        e = OxmlElement(f"w:{side}")
        e.set(qn("w:val"), val)
        if sz is not None:
            e.set(qn("w:sz"), sz)
        e.set(qn("w:color"), "auto")
        e.set(qn("w:space"), "0")
        return e

    for tbl in doc.tables:
        # 跳过封面表格（第一行第一个 cell 含"论文题目"），保留其下划线样式
        if len(tbl.rows) > 0 and "论文题目" in (tbl.rows[0].cells[0].text or ""):
            continue
        # 跳过代码块表格（单 1x1 cell，文本为等宽代码）——保持无边框
        if _is_code_block_table(tbl):
            _clear_table_borders(tbl)
            continue
        tblPr = tbl._element.find(qn("w:tblPr"))
        if tblPr is None:
            tblPr = OxmlElement("w:tblPr")
            tbl._element.insert(0, tblPr)

        # 1. 整表居中
        jc = tblPr.find(qn("w:jc"))
        if jc is None:
            jc = OxmlElement("w:jc")
            # jc 在 tblPr 里的子元素位置：tblStyle, tblpPr, tblOverlap, bidiVisual, tblStyleRowBandSize,
            # tblStyleColBandSize, tblW, jc, ...
            tblPr.append(jc)
        jc.set(qn("w:val"), "center")

        # 2. 清掉旧边框
        old_tb = tblPr.find(qn("w:tblBorders"))
        if old_tb is not None:
            tblPr.remove(old_tb)

        # 3. 新边框：三线表的上下两条粗线
        tb = OxmlElement("w:tblBorders")
        for side, val, sz in [
            ("top", "single", "4"),
            ("left", "nil", None),
            ("bottom", "single", "4"),
            ("right", "nil", None),
            ("insideH", "nil", None),
            ("insideV", "nil", None),
        ]:
            tb.append(_mkborder(side, val, sz))
        tblPr.append(tb)

        # 4. 清掉所有 tc 的 tcBorders
        for row in tbl.rows:
            for cell in row.cells:
                tcPr = cell._tc.find(qn("w:tcPr"))
                if tcPr is not None:
                    old_tcb = tcPr.find(qn("w:tcBorders"))
                    if old_tcb is not None:
                        tcPr.remove(old_tcb)

        # 5. 首行每 tc 只加 bottom 细线 = 三线表第二条线（标题行/数据行分隔）
        #    不写 top —— 让首行顶部继承 tblBorders 的 top 粗线（即整表顶线 = 第一条线）
        #    末行底部自然继承 tblBorders 的 bottom 粗线 = 第三条线
        #    单行表（rows=1）跳过：首行=末行，不加 header 分隔，避免覆盖整表底线
        if len(tbl.rows) >= 2:
            header_row = tbl.rows[0]
            for cell in header_row.cells:
                tc = cell._tc
                tcPr = tc.find(qn("w:tcPr"))
                if tcPr is None:
                    tcPr = OxmlElement("w:tcPr")
                    tc.insert(0, tcPr)
                tcB = OxmlElement("w:tcBorders")
                # 只写 bottom，其它 side 省略 → 继承 tblBorders
                tcB.append(_mkborder("bottom", "single", "4"))
                tcPr.append(tcB)


def center_all_images(doc) -> None:
    """图片所在段落一律居中（regardless of 样式名）。"""
    for p in doc.paragraphs:
        drawings = p._element.findall(".//" + qn("w:drawing"))
        if drawings:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER


_WP_NS = "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
_A_NS = "http://schemas.openxmlformats.org/drawingml/2006/main"


def resize_all_images_to_width(
    doc,
    width_cm: float = 14.7,
    max_height_cm: float | None = None,
) -> None:
    """所有嵌入图片等比缩放到指定栏宽，高度可选硬上限。

    Word OOXML 里图片尺寸由两组节点决定：
      1. ``<wp:extent cx cy/>`` —— inline/anchor 容器尺寸
      2. ``<a:ext cx cy/>``     —— pic 内部 xfrm 的 transform 尺寸

    两者必须同步缩放，否则图形容器与内部绘图会错位。按 wp:extent 的缩放比
    同步更新 drawing 子树下所有 a:ext，保证容器与 transform 一致。

    若给了 ``max_height_cm``，则最终缩放比取 ``min(width_scale, height_scale)``，
    保证宽≤width_cm 且 高≤max_height_cm，始终等比例不变形。典型用法：
    ``width_cm=14.7, max_height_cm=9.9`` — A4 正文栏宽 + 页高 1/3 (29.7/3) 上限。
    """
    target_cx = int(round(width_cm * 360000))  # 1 cm = 360000 EMU
    max_cy = int(round(max_height_cm * 360000)) if max_height_cm else None
    for drawing in doc.element.iter(qn("w:drawing")):
        wp_extent = None
        for c in drawing.iter(f"{{{_WP_NS}}}extent"):
            wp_extent = c
            break
        if wp_extent is None:
            continue
        old_cx = int(wp_extent.get("cx") or 0)
        old_cy = int(wp_extent.get("cy") or 0)
        if old_cx <= 0 or old_cy <= 0:
            continue
        # 等比缩放：取宽度和高度两个约束的最小缩放比
        scale_w = target_cx / old_cx
        scale = scale_w if max_cy is None else min(scale_w, max_cy / old_cy)
        new_cx = int(round(old_cx * scale))
        new_cy = int(round(old_cy * scale))
        wp_extent.set("cx", str(new_cx))
        wp_extent.set("cy", str(new_cy))
        # 同步 drawing 内所有 a:ext（pic xfrm 的尺寸）
        for a_ext in drawing.iter(f"{{{_A_NS}}}ext"):
            a_cx = int(a_ext.get("cx") or 0)
            a_cy = int(a_ext.get("cy") or 0)
            if a_cx > 0 and a_cy > 0:
                a_ext.set("cx", str(int(round(a_cx * scale))))
                a_ext.set("cy", str(int(round(a_cy * scale))))


def style_code_block_tables(doc) -> None:
    """代码块表格里的段落：五号字（21 half-points = 10.5pt）+ 取消首行缩进 + 左对齐。"""
    for tbl in doc.tables:
        if not _is_code_block_table(tbl):
            continue
        # 遍历 cell 内所有段落
        for tr in tbl._element.findall(qn("w:tr")):
            for tc in tr.findall(qn("w:tc")):
                for p_el in tc.findall(qn("w:p")):
                    pPr = p_el.find(qn("w:pPr"))
                    if pPr is None:
                        pPr = OxmlElement("w:pPr")
                        p_el.insert(0, pPr)
                    # 清首行缩进
                    ind = pPr.find(qn("w:ind"))
                    if ind is None:
                        ind = OxmlElement("w:ind")
                        pPr.append(ind)
                    ind.set(qn("w:firstLine"), "0")
                    ind.set(qn("w:firstLineChars"), "0")
                    ind.set(qn("w:left"), "0")
                    ind.set(qn("w:leftChars"), "0")
                    # 左对齐
                    jc = pPr.find(qn("w:jc"))
                    if jc is None:
                        jc = OxmlElement("w:jc")
                        pPr.append(jc)
                    jc.set(qn("w:val"), "left")
                    # 每个 run 的字号设 21 half-points = 10.5pt (五号)
                    for r in p_el.findall(qn("w:r")):
                        rPr = r.find(qn("w:rPr"))
                        if rPr is None:
                            rPr = OxmlElement("w:rPr")
                            r.insert(0, rPr)
                        sz = rPr.find(qn("w:sz"))
                        if sz is None:
                            sz = OxmlElement("w:sz")
                            rPr.append(sz)
                        sz.set(qn("w:val"), "21")
                        szCs = rPr.find(qn("w:szCs"))
                        if szCs is None:
                            szCs = OxmlElement("w:szCs")
                            rPr.append(szCs)
                        szCs.set(qn("w:val"), "21")


def apply_table_body_font_size(doc) -> None:
    """所有数据型表格 cell 内 run 统一五号字（sz=21 half-points = 10.5pt）。
    跳过：封面表（保留自有排版）/ 代码块表（已在 style_code_block_tables 单独刷五号）。"""
    for tbl in doc.tables:
        if len(tbl.rows) > 0 and "论文题目" in (tbl.rows[0].cells[0].text or ""):
            continue
        if _is_code_block_table(tbl):
            continue
        for tr in tbl._element.findall(qn("w:tr")):
            for tc in tr.findall(qn("w:tc")):
                for p_el in tc.findall(qn("w:p")):
                    for r_el in p_el.findall(qn("w:r")):
                        rPr = r_el.find(qn("w:rPr"))
                        if rPr is None:
                            rPr = OxmlElement("w:rPr")
                            r_el.insert(0, rPr)
                        for tag in ("w:sz", "w:szCs"):
                            e = rPr.find(qn(tag))
                            if e is None:
                                e = OxmlElement(tag)
                                rPr.append(e)
                            e.set(qn("w:val"), "21")


def center_all_table_cells(doc) -> None:
    """所有表格 cell：垂直居中（vAlign=center）+ 水平居中（段 jc=center）。
    抓手：
      1. 跳过封面表（第一行第一 cell 含"论文题目"）—— 保留自有下划线 + 左对齐样式
      2. 跳过代码块表（_is_code_block_table）—— 代码保持左对齐 + 顶对齐
      3. 其余表格（三线表 / 单列表 / 对比表）：cell vAlign=center + 每段 jc=center"""
    for tbl in doc.tables:
        # 封面表保留
        if len(tbl.rows) > 0 and "论文题目" in (tbl.rows[0].cells[0].text or ""):
            continue
        is_code = _is_code_block_table(tbl)
        # 所有 cell：vAlign=center
        for tr in tbl._element.findall(qn("w:tr")):
            for tc in tr.findall(qn("w:tc")):
                tcPr = tc.find(qn("w:tcPr"))
                if tcPr is None:
                    tcPr = OxmlElement("w:tcPr")
                    tc.insert(0, tcPr)
                old_v = tcPr.find(qn("w:vAlign"))
                if old_v is not None:
                    tcPr.remove(old_v)
                vAlign = OxmlElement("w:vAlign")
                vAlign.set(qn("w:val"), "center")
                tcPr.append(vAlign)
                # 代码块 cell：vAlign 中，但段保留左对齐
                if is_code:
                    continue
                # 段 jc=center（水平居中）
                for p_el in tc.findall(qn("w:p")):
                    pPr = p_el.find(qn("w:pPr"))
                    if pPr is None:
                        pPr = OxmlElement("w:pPr")
                        p_el.insert(0, pPr)
                    jc = pPr.find(qn("w:jc"))
                    if jc is None:
                        jc = OxmlElement("w:jc")
                        pPr.append(jc)
                    jc.set(qn("w:val"), "center")
                    # 顺手清首行缩进，避免表格内段首多出 2 字缩进偏移中心
                    ind = pPr.find(qn("w:ind"))
                    if ind is None:
                        ind = OxmlElement("w:ind")
                        pPr.append(ind)
                    ind.set(qn("w:firstLine"), "0")
                    ind.set(qn("w:firstLineChars"), "0")


def add_equation_numbers(doc) -> None:
    """给所有独立公式段（含 <m:oMathPara>）按章内顺序补编号 "(N.M)" 右对齐。
    实现：
    1. 维护当前章号（遇到 Heading 1 非无编号章时更新）
    2. 对每个含 oMathPara 的段：
       - 把 oMathPara 内的 oMath 提出来（去掉 oMathPara 包装）
       - 段内新序列：[tab] + <m:oMath> + [tab] + "(章号.序号)"
       - 段落 pPr 加 tab stops：center tab @ 页中，right tab @ 页右边距"""
    from docx.oxml import OxmlElement as O

    M_NS = "http://schemas.openxmlformats.org/officeDocument/2006/math"
    OMP = f"{{{M_NS}}}oMathPara"
    OM = f"{{{M_NS}}}oMath"

    body = doc.element.body

    # 页面可用宽度（twips）：A4 21cm - 左右边距各 3.17cm = 14.66cm ≈ 8316 twips
    # 1 cm = 566.929 twips
    def _cm2tw(cm: float) -> int:
        return int(round(cm * 566.929))

    avail_cm = 21.0 - 3.17 * 2
    center_pos = _cm2tw(avail_cm / 2)
    right_pos = _cm2tw(avail_cm)

    cur_ch = 0  # 当前章号
    eq_n = 0  # 章内公式序号
    ch_num_re = re.compile(r"^(\d+)\s+")

    for p_el in list(body.iter(qn("w:p"))):
        # 更新章号：遇到 Heading 1 且文字以 "N " 开头
        pPr = p_el.find(qn("w:pPr"))
        if pPr is not None:
            pStyle = pPr.find(qn("w:pStyle"))
            if pStyle is not None and pStyle.get(qn("w:val")) in ("Heading1", "1"):
                # 读段文字
                txt = "".join(t.text or "" for t in p_el.iter(qn("w:t")))
                m = ch_num_re.match(txt.strip())
                if m:
                    new_ch = int(m.group(1))
                    if new_ch != cur_ch:
                        cur_ch = new_ch
                        eq_n = 0

        # 找段内的 oMathPara（0 或 1 个，多数情况下 1 个）
        omp_list = p_el.findall(".//" + OMP)
        if not omp_list or cur_ch == 0:
            continue
        eq_n += 1
        label = f"（{cur_ch}-{eq_n}）"  # 规范示例 "1－1"，采用中文全角括号 + 半角短横

        omp = omp_list[0]
        # 取 oMathPara 里的 oMath
        omath = omp.find(OM)
        if omath is None:
            continue

        # 在 p_el 里定位 omp，用新序列替换
        # 新序列：<w:r><w:tab/></w:r>  <m:oMath>...</m:oMath>  <w:r><w:tab/></w:r>  <w:r><w:t>(N.M)</w:t></w:r>
        parent = omp.getparent()
        idx = list(parent).index(omp)
        parent.remove(omp)

        def _mk_tab_run():
            r = O("w:r")
            tab = O("w:tab")
            r.append(tab)
            return r

        def _mk_label_run(text):
            r = O("w:r")
            rPr = O("w:rPr")
            rF = O("w:rFonts")
            rF.set(qn("w:ascii"), "Times New Roman")
            rF.set(qn("w:hAnsi"), "Times New Roman")
            rF.set(qn("w:eastAsia"), "宋体")
            rPr.append(rF)
            sz = O("w:sz")
            sz.set(qn("w:val"), "24")
            rPr.append(sz)
            szCs = O("w:szCs")
            szCs.set(qn("w:val"), "24")
            rPr.append(szCs)
            r.append(rPr)
            t = O("w:t")
            t.text = text
            t.set(qn("xml:space"), "preserve")
            r.append(t)
            return r

        new_nodes = [
            _mk_tab_run(),  # tab 1 → center
            omath,
            _mk_tab_run(),  # tab 2 → right
            _mk_label_run(label),
        ]
        for offset, node in enumerate(new_nodes):
            parent.insert(idx + offset, node)

        # 段落 pPr：添加 tab stops（center + right）+ 左对齐
        if pPr is None:
            pPr = O("w:pPr")
            p_el.insert(0, pPr)
        # 清旧 tabs
        old_tabs = pPr.find(qn("w:tabs"))
        if old_tabs is not None:
            pPr.remove(old_tabs)
        tabs = O("w:tabs")
        t_c = O("w:tab")
        t_c.set(qn("w:val"), "center")
        t_c.set(qn("w:pos"), str(center_pos))
        tabs.append(t_c)
        t_r = O("w:tab")
        t_r.set(qn("w:val"), "right")
        t_r.set(qn("w:pos"), str(right_pos))
        tabs.append(t_r)
        pPr.append(tabs)
        # 段落左对齐（tab 才能生效）
        jc = pPr.find(qn("w:jc"))
        if jc is None:
            jc = O("w:jc")
            pPr.append(jc)
        jc.set(qn("w:val"), "left")


def wrap_listings_and_algorithms(doc) -> None:
    """代码清单（SourceCode）+ 算法伪代码段 → 单行单格 + 全外框表格。
    抓手：
      1. 连续 SourceCode 段分组，整组塞进一个单格表格
      2. 算法伪代码 BodyText 段（pandoc 压扁后的残留，文本多 "//" + ForEach/排序/赋值箭头特征）单独塞进一个单格表格
      3. 表格 tblBorders 全 single sz=4：顶/左/底/右 + insideH/insideV
      4. 表格本身排除在三线表 pass 之外 —— 所以本函数必须在 apply_three_line_tables 之后跑
         （否则三线表 pass 会把我们新建的全外框表也降级成三线表）"""
    import re as _re

    def _new_framed_table():
        tbl = OxmlElement("w:tbl")
        # tblPr
        tblPr = OxmlElement("w:tblPr")
        tblW = OxmlElement("w:tblW")
        tblW.set(qn("w:type"), "pct")
        tblW.set(qn("w:w"), "5000")  # 100%
        tblPr.append(tblW)
        jc = OxmlElement("w:jc")
        jc.set(qn("w:val"), "center")
        tblPr.append(jc)
        # 全外框 tblBorders
        tb = OxmlElement("w:tblBorders")
        for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
            b = OxmlElement(f"w:{side}")
            b.set(qn("w:val"), "single")
            b.set(qn("w:sz"), "4")
            b.set(qn("w:color"), "auto")
            b.set(qn("w:space"), "0")
            tb.append(b)
        tblPr.append(tb)
        tbl.append(tblPr)
        # tblGrid 单列
        tg = OxmlElement("w:tblGrid")
        gc = OxmlElement("w:gridCol")
        gc.set(qn("w:w"), "8311")  # A4 - 左右边距 ≈ 14.66cm twips
        tg.append(gc)
        tbl.append(tg)
        return tbl

    def _new_tr_tc():
        tr = OxmlElement("w:tr")
        tc = OxmlElement("w:tc")
        tcPr = OxmlElement("w:tcPr")
        tcW = OxmlElement("w:tcW")
        tcW.set(qn("w:type"), "pct")
        tcW.set(qn("w:w"), "5000")
        tcPr.append(tcW)
        tc.append(tcPr)
        tr.append(tc)
        return tr, tc

    body = doc.element.body

    # ---------- 1. 找 SourceCode 连续段组 ----------
    groups = []  # list of list[p_el]
    cur = []
    # 只遍历 body 直接子 p（避免深入已有表格内）
    for child in list(body):
        if child.tag == qn("w:p"):
            pPr = child.find(qn("w:pPr"))
            sid = None
            if pPr is not None:
                ps = pPr.find(qn("w:pStyle"))
                if ps is not None:
                    sid = ps.get(qn("w:val"))
            if sid == "SourceCode":
                cur.append(child)
                continue
        # 不是 SourceCode 段 → 关闭当前组
        if cur:
            groups.append(cur)
            cur = []
    if cur:
        groups.append(cur)

    # ---------- 2. 找 algorithm 伪代码段 ----------
    # flatten.rewrite_algorithms 现在把算法体写成 \begin{quote}...\end{quote}，
    # pandoc 落 Block Text 样式，每行独立段落，行内 $...$ 渲染为 OMML 公式。
    # 识别模式：一个 "算法 X.Y ..." 粗体 caption 段（Body Text）之后，
    # 紧跟若干 Block Text 段——整段连续区间归为一组，wrap 成单格全框表格。
    _ALGO_CAP_RE = _re.compile(r"^算法\s*\d+[\.-]\d+\s")

    def _pstyle(el) -> str | None:
        pPr = el.find(qn("w:pPr"))
        if pPr is None:
            return None
        ps = pPr.find(qn("w:pStyle"))
        return ps.get(qn("w:val")) if ps is not None else None

    def _all_runs_bold(el) -> bool:
        runs = el.findall(qn("w:r"))
        has_text_run = False
        for r in runs:
            t = "".join((tt.text or "") for tt in r.findall(qn("w:t")))
            if not t.strip():
                continue
            has_text_run = True
            rPr = r.find(qn("w:rPr"))
            if rPr is None or rPr.find(qn("w:b")) is None:
                return False
        return has_text_run

    algo_groups = []
    children_list = list(body)
    i = 0
    while i < len(children_list):
        child = children_list[i]
        if child.tag != qn("w:p"):
            i += 1
            continue
        txt = "".join(t.text or "" for t in child.iter(qn("w:t"))).lstrip()
        if not _ALGO_CAP_RE.match(txt) or not _all_runs_bold(child):
            i += 1
            continue
        # caption 命中：收集紧随其后的 Block Text 段
        group: list = []
        j = i + 1
        while j < len(children_list):
            nxt = children_list[j]
            if nxt.tag != qn("w:p"):
                break
            if _pstyle(nxt) not in ("BlockText", "Block Text"):
                break
            group.append(nxt)
            j += 1
        if group:
            algo_groups.append(group)
        i = j

    # ---------- 3. 每组段用单格全外框表格替换 ----------
    for group in groups + algo_groups:
        first = group[0]
        parent = first.getparent()
        idx = list(parent).index(first)

        tbl = _new_framed_table()
        tr, tc = _new_tr_tc()
        tbl.append(tr)

        # 把段 move 到 tc 里
        for p_el in group:
            p_el.getparent().remove(p_el)
            tc.append(p_el)

        parent.insert(idx, tbl)


def flatten_list_indent(doc) -> None:
    """numbering.xml 的所有 lvl 默认 left=720 hanging=360，导致 pandoc 转出的
    enumerate/itemize 列表段编号左侧有 1.27cm 缩进（规范要求数字编号顶格）。
    抓手：遍历 numbering.xml 的所有 <w:lvl>/<w:pPr>/<w:ind>，left/hanging 清零。"""
    # --- 1. 修改 numbering.xml 根元素（python-docx 会在 save 时自动 serialize） ---
    numbering_part = None
    for rel in doc.part.rels.values():
        if rel.reltype.endswith("/numbering"):
            numbering_part = rel.target_part
            break
    if numbering_part is not None and hasattr(numbering_part, "element"):
        root = numbering_part.element  # CT_Numbering (lxml wrapper)
        W = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
        for lvl in root.iter(W + "lvl"):
            # 编号格式统一为（1）（2）…（规范要求中文括号包裹阿拉伯数字）
            ilvl_val = lvl.get(W + "ilvl", "0")
            numFmt_el = lvl.find(W + "numFmt")
            if numFmt_el is not None and numFmt_el.get(W + "val", "") == "decimal":
                lvlText_el = lvl.find(W + "lvlText")
                if lvlText_el is not None and ilvl_val == "0":
                    lvlText_el.set(W + "val", "（%1）")

            pPr = lvl.find(W + "pPr")
            if pPr is None:
                continue
            ind = pPr.find(W + "ind")
            if ind is None:
                # 显式新建 ind=0（防 Word 套老默认）
                from lxml import etree as _et

                ind = _et.SubElement(pPr, W + "ind")
            # 清 left/leftChars/hanging/hangingChars/firstLine/firstLineChars
            for attr in (
                "left",
                "leftChars",
                "hanging",
                "hangingChars",
                "firstLine",
                "firstLineChars",
            ):
                key = W + attr
                if key in ind.attrib:
                    del ind.attrib[key]
            # 编号行首行缩进2字符
            ind.set(W + "firstLineChars", "200")

    # --- 2. 段级 ind 清零（双保险，且覆盖 numbering.xml 继承） ---
    for p in doc.paragraphs:
        pPr = p._element.find(qn("w:pPr"))
        if pPr is None:
            continue
        numPr = pPr.find(qn("w:numPr"))
        if numPr is None:
            continue
        ind = pPr.find(qn("w:ind"))
        if ind is None:
            ind = OxmlElement("w:ind")
            pPr.append(ind)
        for attr in (
            "left",
            "leftChars",
            "hanging",
            "hangingChars",
            "firstLine",
            "firstLineChars",
        ):
            key = qn(f"w:{attr}")
            if key in ind.attrib:
                del ind.attrib[key]
        # 编号行首行缩进2字符，与正文段对齐
        ind.set(qn("w:firstLineChars"), "200")
