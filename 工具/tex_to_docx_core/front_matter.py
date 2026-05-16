from __future__ import annotations

import re
import sys
from copy import deepcopy

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from .config import (
    COVER_LABEL_COL_W,
    COVER_VALUE_COL_W,
    TEMPLATE_DOCX,
    TEMPLATE_PREAMBLE_RANGE,
    THESIS_DIR,
)
from .docx_common import _set_rfonts
from .docx_structure import _make_toc_block


def _parse_info_tex() -> dict[str, str]:
    """从 info.tex 提取所有 \\newcommand 宏值。"""
    info = THESIS_DIR / "info.tex"
    if not info.exists():
        return {}
    text = info.read_text(encoding="utf-8")
    result = {}
    for m in re.finditer(r"\\newcommand\{\\(\w+)\}\{([^}]*)\}", text):
        result[m.group(1)] = m.group(2)
    return result


def _reset_cover_table_borders(tbl) -> None:
    """封面表格边框闭环：清掉 table 级外框与内框，仅给第二列每行加下边框。
    同时调整列宽与居中：第一列约 4em、第二列约 8cm，整表居中 + fixed layout。"""
    from docx.oxml import OxmlElement

    t_el = tbl._element
    # 1. 表级 tblBorders 全清为 nil
    tblPr = t_el.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        t_el.insert(0, tblPr)
    # 去掉模板自带表样式，避免字号/边距/加粗等继承污染封面
    old_style = tblPr.find(qn("w:tblStyle"))
    if old_style is not None:
        tblPr.remove(old_style)
    tblBorders = tblPr.find(qn("w:tblBorders"))
    if tblBorders is not None:
        tblPr.remove(tblBorders)
    tblBorders = OxmlElement("w:tblBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"), "nil")
        tblBorders.append(b)
    tblPr.append(tblBorders)
    tblCellMar = tblPr.find(qn("w:tblCellMar"))
    if tblCellMar is not None:
        tblPr.remove(tblCellMar)
    tblCellMar = OxmlElement("w:tblCellMar")
    for side in ("top", "left", "bottom", "right"):
        m = OxmlElement(f"w:{side}")
        m.set(qn("w:w"), "0")
        m.set(qn("w:type"), "dxa")
        tblCellMar.append(m)
    tblPr.append(tblCellMar)

    # 1.5 列宽 / 整表居中 / fixed layout 防自动压缩
    COL0_W = COVER_LABEL_COL_W
    COL1_W = COVER_VALUE_COL_W
    TOTAL_W = COL0_W + COL1_W

    # tblW: 总宽
    tblW = tblPr.find(qn("w:tblW"))
    if tblW is None:
        tblW = OxmlElement("w:tblW")
        tblPr.append(tblW)
    tblW.set(qn("w:w"), str(TOTAL_W))
    tblW.set(qn("w:type"), "dxa")

    # jc: 居中
    jc = tblPr.find(qn("w:jc"))
    if jc is None:
        jc = OxmlElement("w:jc")
        tblPr.append(jc)
    jc.set(qn("w:val"), "center")

    # tblLayout: fixed，防止 Word 根据内容自动重排列宽
    tblLayout = tblPr.find(qn("w:tblLayout"))
    if tblLayout is None:
        tblLayout = OxmlElement("w:tblLayout")
        tblPr.append(tblLayout)
    tblLayout.set(qn("w:type"), "fixed")

    # tblGrid
    tblGrid = t_el.find(qn("w:tblGrid"))
    if tblGrid is not None:
        t_el.remove(tblGrid)
    tblGrid = OxmlElement("w:tblGrid")
    for w in (COL0_W, COL1_W):
        gc = OxmlElement("w:gridCol")
        gc.set(qn("w:w"), str(w))
        tblGrid.append(gc)
    # tblGrid 必须放在 tblPr 之后、第一个 tr 之前
    tblPr.addnext(tblGrid)

    # 每行 tc 的 tcW 也重写
    for row in tbl.rows:
        tcs = row._element.findall(qn("w:tc"))
        for ci, tc in enumerate(tcs):
            tcPr = tc.find(qn("w:tcPr"))
            if tcPr is None:
                tcPr = OxmlElement("w:tcPr")
                tc.insert(0, tcPr)
            tcW = tcPr.find(qn("w:tcW"))
            if tcW is None:
                tcW = OxmlElement("w:tcW")
                tcPr.insert(0, tcW)
            tcW.set(qn("w:w"), str(COL0_W if ci == 0 else COL1_W))
            tcW.set(qn("w:type"), "dxa")

    # 2. 每个 cell：先清 tcBorders，然后给第二列每行（含标题 2 行）加 bottom single
    for ri, row in enumerate(tbl.rows):
        tcs = row._element.findall(qn("w:tc"))
        for ci, tc in enumerate(tcs):
            tcPr = tc.find(qn("w:tcPr"))
            if tcPr is None:
                tcPr = OxmlElement("w:tcPr")
                tc.insert(0, tcPr)
            # 移除现有 tcBorders
            old_tcBorders = tcPr.find(qn("w:tcBorders"))
            if old_tcBorders is not None:
                tcPr.remove(old_tcBorders)
            # 新 tcBorders：全 nil，第二列底部 single
            tcBorders = OxmlElement("w:tcBorders")
            for side in ("top", "left", "right"):
                b = OxmlElement(f"w:{side}")
                b.set(qn("w:val"), "nil")
                tcBorders.append(b)
            bottom = OxmlElement("w:bottom")
            if ci == 1:  # 第二列
                bottom.set(qn("w:val"), "single")
                bottom.set(qn("w:sz"), "4")
                bottom.set(qn("w:color"), "auto")
            else:
                bottom.set(qn("w:val"), "nil")
            tcBorders.append(bottom)
            tcPr.append(tcBorders)


def fill_cover_info(doc) -> None:
    """填充封面表格（模板搬来的 8行x2列表格）和日期行。
    复刻 tex 原版效果：左列标签右对齐无冒号，右列内容居中 + 下划线。"""
    from docx.shared import Pt

    info = _parse_info_tex()
    if not info:
        return

    # 封面表格行映射：行号 → 值（表格第2列）
    title_a = info.get("thesisTitleZhA", "")
    title_b = info.get("thesisTitleZhB", "")

    cell_values = {
        0: title_a,
        1: title_b,
        2: info.get("thesisCollege", ""),
        3: info.get("thesisMajor", ""),
        4: info.get("thesisClass", ""),
        5: info.get("thesisStudentId", ""),
        6: info.get("thesisAuthor", ""),
        7: info.get("thesisAdvisor", ""),
    }

    label_values = {
        0: "论文题目：",
        1: "",
        2: "学    院：",
        3: "专    业：",
        4: "班    级：",
        5: "学    号：",
        6: "姓    名：",
        7: "指导教师：",
    }

    # 找封面表格（第一个含"论文题目"的表格）
    cover_tbl = None
    for t in doc.tables:
        for row in t.rows[:1]:
            if "论文题目" in (row.cells[0].text or ""):
                cover_tbl = t
                break
        if cover_tbl:
            break

    if cover_tbl:
        for row_idx in range(min(8, len(cover_tbl.rows))):
            row = cover_tbl.rows[row_idx]
            tcs = row._element.findall(qn("w:tc"))

            # --- 左列：重写标签 + 右对齐 ---
            tc0 = tcs[0]
            p0 = tc0.find(qn("w:p"))
            if p0 is None:
                p0 = OxmlElement("w:p")
                tc0.append(p0)
            # 清空段内所有 run
            for r in list(p0.findall(qn("w:r"))):
                p0.remove(r)
            # 设 jc=right
            pPr0 = p0.find(qn("w:pPr"))
            if pPr0 is None:
                pPr0 = OxmlElement("w:pPr")
                p0.insert(0, pPr0)
            # 去掉模板残留的段内粗体/字号继承，避免封面文字过重过宽
            old_rPr0 = pPr0.find(qn("w:rPr"))
            if old_rPr0 is not None:
                pPr0.remove(old_rPr0)
            jc0 = pPr0.find(qn("w:jc"))
            if jc0 is None:
                jc0 = OxmlElement("w:jc")
                pPr0.append(jc0)
            jc0.set(qn("w:val"), "right")
            # 清首行缩进
            ind0 = pPr0.find(qn("w:ind"))
            if ind0 is not None:
                pPr0.remove(ind0)
            # 写标签 run
            label = label_values.get(row_idx, "")
            if label:
                r_el = OxmlElement("w:r")
                rPr = OxmlElement("w:rPr")
                sz = OxmlElement("w:sz")
                sz.set(qn("w:val"), "30")
                rPr.append(sz)
                szCs = OxmlElement("w:szCs")
                szCs.set(qn("w:val"), "30")
                rPr.append(szCs)
                rF = OxmlElement("w:rFonts")
                rF.set(qn("w:ascii"), "Times New Roman")
                rF.set(qn("w:hAnsi"), "Times New Roman")
                rF.set(qn("w:eastAsia"), "宋体")
                rPr.append(rF)
                if row_idx == 0:
                    rPr.append(OxmlElement("w:b"))
                    rPr.append(OxmlElement("w:bCs"))
                r_el.append(rPr)
                t_el = OxmlElement("w:t")
                t_el.text = label
                t_el.set(qn("xml:space"), "preserve")
                r_el.append(t_el)
                p0.append(r_el)

            # --- 右列：重写内容 + 居中 ---
            tc1 = tcs[1]
            p1 = tc1.find(qn("w:p"))
            if p1 is None:
                p1 = OxmlElement("w:p")
                tc1.append(p1)
            for r in list(p1.findall(qn("w:r"))):
                p1.remove(r)
            # 设 jc=center
            pPr1 = p1.find(qn("w:pPr"))
            if pPr1 is None:
                pPr1 = OxmlElement("w:pPr")
                p1.insert(0, pPr1)
            old_rPr1 = pPr1.find(qn("w:rPr"))
            if old_rPr1 is not None:
                pPr1.remove(old_rPr1)
            jc1 = pPr1.find(qn("w:jc"))
            if jc1 is None:
                jc1 = OxmlElement("w:jc")
                pPr1.append(jc1)
            jc1.set(qn("w:val"), "center")
            # 清首行缩进
            ind1 = pPr1.find(qn("w:ind"))
            if ind1 is not None:
                pPr1.remove(ind1)
            # 写内容 run
            value = cell_values.get(row_idx, "")
            if value:
                r_el = OxmlElement("w:r")
                rPr = OxmlElement("w:rPr")
                sz = OxmlElement("w:sz")
                sz.set(qn("w:val"), "30")
                rPr.append(sz)
                szCs = OxmlElement("w:szCs")
                szCs.set(qn("w:val"), "30")
                rPr.append(szCs)
                rF = OxmlElement("w:rFonts")
                rF.set(qn("w:ascii"), "Times New Roman")
                rF.set(qn("w:hAnsi"), "Times New Roman")
                rF.set(qn("w:eastAsia"), "宋体")
                rPr.append(rF)
                if row_idx <= 1:
                    rPr.append(OxmlElement("w:b"))
                    rPr.append(OxmlElement("w:bCs"))
                r_el.append(rPr)
                t_el = OxmlElement("w:t")
                t_el.text = value
                t_el.set(qn("xml:space"), "preserve")
                r_el.append(t_el)
                p1.append(r_el)

        _reset_cover_table_borders(cover_tbl)

    # 更新日期段文本
    for p in doc.paragraphs[:30]:
        txt = (p.text or "").strip()
        if "年" in txt and "月" in txt and ("日" in txt or txt.endswith("月")):
            date_text = info.get("thesisDate", "2026年5月")
            # 解析年月，格式化为 "年     月     日"
            m = re.match(r"(\d{4})年(\d{1,2})月(?:(\d{1,2})日)?", date_text)
            if m:
                date_text = f"{m.group(1)}年     {m.group(2)}月     {m.group(3) or ''}日"
            for r in list(p.runs):
                r._element.getparent().remove(r._element)
            run = p.add_run(date_text)
            run.font.size = Pt(15)
            _set_rfonts(run, ascii_="Times New Roman", cjk="宋体")
            break

    # 填充"所属学院：  所属专业："行（声明+授权页各一处）
    college = info.get("thesisCollege", "")
    major = info.get("thesisMajor", "")
    for p in doc.paragraphs[:60]:
        txt = p.text or ""
        if "所属学院" in txt and "所属专业" in txt:
            for r in list(p.runs):
                r._element.getparent().remove(r._element)
            new_text = f"所属学院：{college}     所属专业：{major}"
            run = p.add_run(new_text)
            run.font.size = Pt(12)
            _set_rfonts(run, ascii_="Times New Roman", cjk="宋体")

    # 填充"毕业论文（设计）名称："行（声明页 + 授权页各一处）
    title_full = (
        info.get("thesisTitleZhA", "") + info.get("thesisTitleZhB", "")
    ).strip()
    for p in doc.paragraphs[:60]:
        txt = (p.text or "").strip()
        if txt == "毕业论文（设计）名称：" or txt.rstrip() == "毕业论文（设计）名称：":
            for r in list(p.runs):
                r._element.getparent().remove(r._element)
            run = p.add_run(f"毕业论文（设计）名称：{title_full}")
            run.font.size = Pt(12)
            _set_rfonts(run, ascii_="Times New Roman", cjk="宋体")

    # 修正签名行：用制表符（tab stops）排版，代替空格硬拉
    # 底层逻辑：空格在不同字号/字体下对不齐，tab stops 绝对定位才是正解
    # 布局：签名前缀 [TAB→pos1] 日期： [TAB→pos2]      年   月   日
    sign_prefixes = ["论文作者（签名）：", "授权人（学生签名）：", "指导教师（签名）："]
    for p in doc.paragraphs[:60]:
        txt = p.text or ""
        for pref in sign_prefixes:
            if txt.startswith(pref) and "日期" in txt:
                p_el = p._element
                # 清空 run
                for r in list(p_el.findall(qn("w:r"))):
                    p_el.remove(r)
                # 确保有 pPr
                pPr = p_el.find(qn("w:pPr"))
                if pPr is None:
                    pPr = OxmlElement("w:pPr")
                    p_el.insert(0, pPr)
                # 移除旧 tabs
                old_tabs = pPr.find(qn("w:tabs"))
                if old_tabs is not None:
                    pPr.remove(old_tabs)
                # 新 tabs：2 个 left tab
                tabs = OxmlElement("w:tabs")
                for pos in ("4500", "5800"):
                    tab = OxmlElement("w:tab")
                    tab.set(qn("w:val"), "left")
                    tab.set(qn("w:pos"), pos)
                    tabs.append(tab)
                pPr.append(tabs)

                # 构造 run 序列
                def _mk_run(text):
                    r_el = OxmlElement("w:r")
                    rPr = OxmlElement("w:rPr")
                    sz = OxmlElement("w:sz")
                    sz.set(qn("w:val"), "24")
                    rPr.append(sz)
                    szCs = OxmlElement("w:szCs")
                    szCs.set(qn("w:val"), "24")
                    rPr.append(szCs)
                    rF = OxmlElement("w:rFonts")
                    rF.set(qn("w:ascii"), "Times New Roman")
                    rF.set(qn("w:hAnsi"), "Times New Roman")
                    rF.set(qn("w:eastAsia"), "宋体")
                    rPr.append(rF)
                    r_el.append(rPr)
                    t_el = OxmlElement("w:t")
                    t_el.text = text
                    t_el.set(qn("xml:space"), "preserve")
                    r_el.append(t_el)
                    return r_el

                def _mk_tab_run():
                    r_el = OxmlElement("w:r")
                    rPr = OxmlElement("w:rPr")
                    sz = OxmlElement("w:sz")
                    sz.set(qn("w:val"), "24")
                    rPr.append(sz)
                    szCs = OxmlElement("w:szCs")
                    szCs.set(qn("w:val"), "24")
                    rPr.append(szCs)
                    rF = OxmlElement("w:rFonts")
                    rF.set(qn("w:ascii"), "Times New Roman")
                    rF.set(qn("w:hAnsi"), "Times New Roman")
                    rF.set(qn("w:eastAsia"), "宋体")
                    rPr.append(rF)
                    r_el.append(rPr)
                    tab_el = OxmlElement("w:tab")
                    r_el.append(tab_el)
                    return r_el

                # 序列：前缀 + [TAB] + "日期：" + [TAB] + "年   月   日"
                p_el.append(_mk_run(pref))
                p_el.append(_mk_tab_run())
                p_el.append(_mk_run("日期："))
                p_el.append(_mk_tab_run())
                p_el.append(_mk_run("年   月   日"))
                break

    # "目  录"段前的分页由 setup_page_numbers_and_sections 的 sectPr(nextPage) 处理，
    # 不再在此重复插入分页符。

    # 压缩封面到一页：清掉冗余空段，用段前间距控制视觉留白
    _compact_cover_page(doc)


def _compact_cover_page(doc) -> None:
    """把封面 block（body 开头到第一个 sectPr 前）压成一页：
    - 删除所有空段
    - 保留：学校名、"毕业论文"大标题、封面表格、日期段
    - 用段前/段后间距控制留白，防止跨页。"""
    from docx.oxml import OxmlElement

    body = doc.element.body

    # 找封面区结束位置：第一个 sectPr
    cover_end = None
    for i, child in enumerate(body):
        if child.tag == qn("w:p"):
            pPr = child.find(qn("w:pPr"))
            if pPr is not None and pPr.find(qn("w:sectPr")) is not None:
                cover_end = i
                break
        elif child.tag == qn("w:sectPr"):
            cover_end = i
            break
    if cover_end is None:
        return

    # 扫描 [0..cover_end)，找出空段并删除
    to_remove = []
    children = list(body)[:cover_end]
    for i, child in enumerate(children):
        if child.tag == qn("w:p"):
            t_els = child.findall(".//" + qn("w:t"))
            txt = "".join(t.text or "" for t in t_els).strip()
            if not txt:
                # 空段但要保留持有 sectPr 的那个
                pPr = child.find(qn("w:pPr"))
                if pPr is not None and pPr.find(qn("w:sectPr")) is not None:
                    continue
                to_remove.append(child)

    for el in to_remove:
        body.remove(el)

    # 重新扫描取引用
    def _find_para_by_text(keyword_check):
        for i, c in enumerate(list(body)):
            if c.tag == qn("w:p"):
                t_els = c.findall(".//" + qn("w:t"))
                txt = "".join(t.text or "" for t in t_els).strip()
                if keyword_check(txt):
                    return c
        return None

    school_p = _find_para_by_text(lambda t: "湖北文理学院" in t and len(t) < 15)
    title_p = _find_para_by_text(
        lambda t: "毕" in t and "业" in t and "论" in t and "文" in t and len(t) < 20
    )
    date_p = _find_para_by_text(lambda t: "年" in t and "月" in t and len(t) < 30)

    def _set_spacing_before(p_el, before_twips: int):
        if p_el is None:
            return
        pPr = p_el.find(qn("w:pPr"))
        if pPr is None:
            pPr = OxmlElement("w:pPr")
            p_el.insert(0, pPr)
        spacing = pPr.find(qn("w:spacing"))
        if spacing is None:
            spacing = OxmlElement("w:spacing")
            pPr.append(spacing)
        spacing.set(qn("w:before"), str(before_twips))
        spacing.set(qn("w:beforeLines"), "0")

    # 排版：
    # 顶部 → 学校名（段前 600 ≈ 1cm）
    # 学校名 → 毕业论文大标题（段前 1200 ≈ 2cm）
    # 毕业论文 → 表格：用段前空隙给表格放缓冲（段前 1500 ≈ 2.6cm）
    # 表格 → 日期（段前 1500 ≈ 2.6cm）
    _set_spacing_before(school_p, 600)
    _set_spacing_before(title_p, 1200)
    _set_spacing_before(date_p, 1500)


def prepend_front_matter(doc) -> None:
    """把模板 docx 的封面+声明+授权整块搬到 thesis 开头，再插入 TOC 目录域。"""
    if not TEMPLATE_DOCX.exists():
        print(
            f"[warn] 模板 docx 不存在：{TEMPLATE_DOCX}，跳过前置页搬运", file=sys.stderr
        )
        return

    tpl = Document(str(TEMPLATE_DOCX))
    start, end = TEMPLATE_PREAMBLE_RANGE

    # 深拷贝模板段 XML，但剥离会指向错误 rId 的引用节点：
    # - w:headerReference / w:footerReference：rId 在模板里指 header/footer，
    #   在 thesis 里 rId 空间不同，直接拷贝会 Word "无法读取" 报错。
    # - w:pStyle：模板样式名在 thesis.docx 的 styles.xml 不存在，会回退 Normal。保留 OK。
    STRIP_TAGS = {
        qn("w:headerReference"),
        qn("w:footerReference"),
    }

    def _strip_bad_refs(el):
        for bad in list(el.iter()):
            if bad.tag in STRIP_TAGS:
                parent = bad.getparent()
                if parent is not None:
                    parent.remove(bad)

    preamble_elems = []
    tpl_body = tpl.element.body
    tpl_children = list(tpl_body)
    # 取 body 子元素 [start..end]（含 w:p 和 w:tbl），跳过末尾 sectPr
    for i in range(start, min(end + 1, len(tpl_children))):
        child = tpl_children[i]
        if child.tag == qn("w:sectPr"):
            continue
        el = deepcopy(child)
        _strip_bad_refs(el)
        preamble_elems.append(el)

    body = doc.element.body
    first_p = doc.paragraphs[0]._element
    insert_idx = list(body).index(first_p)

    for el in reversed(preamble_elems):
        body.insert(insert_idx, el)
    insert_idx = list(body).index(first_p)
    for el in reversed(_make_toc_block(doc)):
        body.insert(insert_idx, el)


def compress_declaration_and_authorization(doc) -> None:
    """确保原创性声明从新页开始，且声明和授权书在同一页（keepLines+keepNext）。"""
    start_idx = end_idx = None
    for i, p in enumerate(doc.paragraphs[:80]):
        txt = (p.text or "").strip()
        if ("原创性声明" in txt or "独创性声明" in txt) and start_idx is None:
            start_idx = i
            # 声明从新页开始
            pPr = p._element.find(qn("w:pPr"))
            if pPr is None:
                pPr = OxmlElement("w:pPr")
                p._element.insert(0, pPr)
            if pPr.find(qn("w:pageBreakBefore")) is None:
                pPr.append(OxmlElement("w:pageBreakBefore"))
        if "版权使用授权书" in txt:
            pPr = p._element.find(qn("w:pPr"))
            if pPr is not None:
                old = pPr.find(qn("w:pageBreakBefore"))
                if old is not None:
                    pPr.remove(old)
        if txt.startswith("指导教师") and "签名" in txt:
            end_idx = i
            break
    # 给声明+授权区域所有段落加 keepNext + 清除段后间距，防止跨页
    if start_idx is not None and end_idx is not None:
        for p in doc.paragraphs[start_idx : end_idx + 1]:
            pPr = p._element.find(qn("w:pPr"))
            if pPr is None:
                pPr = OxmlElement("w:pPr")
                p._element.insert(0, pPr)
            if pPr.find(qn("w:keepNext")) is None:
                pPr.append(OxmlElement("w:keepNext"))
            spacing = pPr.find(qn("w:spacing"))
            if spacing is None:
                spacing = OxmlElement("w:spacing")
                pPr.append(spacing)
            spacing.set(qn("w:after"), "0")


def tab_align_cover_and_signature_rows(doc) -> None:
    """封面 + 声明/授权签名页里"左标签 + 一堆空格 + 右字段"类行：
    - "2026届本科生毕业论文...存档编号："
    - "论文作者（签名）：...日期：..."
    - "授权人（学生签名）：...日期：..."
    - "指导教师（签名）：...日期：..."
    抓手：
    1. 按正则劈分左右两半（"标签 + 任意空格/冒号 + 后续字段"）
    2. 清空原段所有 run，重建：左 run + <w:tab/> + 右 run
    3. 段落 pPr 追加右对齐 tab stop（pos = A4 可用宽度 = 8311 twips ≈ 14.66cm）"""
    import re as _re

    # (正则匹配段首字符, 劈分点正则) —— 劈分点之前是左半，之后（含）是右半
    SPLIT_RULES = [
        # 2026届本科生毕业论文                存档编号：  →  左="2026届..." 右="存档编号："
        (_re.compile(r"届.*毕业论文"), _re.compile(r"\s{2,}(存档编号[:：].*)$")),
        # 所属学院：XX     所属专业：YY
        (_re.compile(r"^所属学院"), _re.compile(r"\s{2}(所属专业[:：].*)$")),
    ]

    # A4 可用宽度：21 - 3.17*2 = 14.66cm ≈ 8311 twips
    right_pos = 8311

    for p in doc.paragraphs[:60]:  # 前置页只在前 60 段
        txt = p.text or ""
        for prefix_re, split_re in SPLIT_RULES:
            if not prefix_re.search(txt):
                continue
            m = split_re.search(txt)
            if not m:
                continue
            left_part = txt[: m.start()].rstrip()
            right_part = m.group(1).rstrip()

            # 1. 清空原段所有 run（保留 pPr）
            for r in list(p._element.findall(qn("w:r"))):
                p._element.remove(r)

            # 2. 新建左 run + tab + 右 run（字体显式写 TNR / 宋体，防字体正则化乱动）
            def _mk_run(text: str):
                r = OxmlElement("w:r")
                rPr = OxmlElement("w:rPr")
                rF = OxmlElement("w:rFonts")
                rF.set(qn("w:ascii"), "Times New Roman")
                rF.set(qn("w:hAnsi"), "Times New Roman")
                rF.set(qn("w:eastAsia"), "宋体")
                rPr.append(rF)
                r.append(rPr)
                t = OxmlElement("w:t")
                t.text = text
                t.set(qn("xml:space"), "preserve")
                r.append(t)
                return r

            tab_run = OxmlElement("w:r")
            tab_run.append(OxmlElement("w:tab"))

            # 顺序：左 run → tab run → 右 run
            p._element.append(_mk_run(left_part))
            p._element.append(tab_run)
            p._element.append(_mk_run(right_part))

            # 3. 段落 pPr 加右对齐 tab stop
            pPr = p._element.find(qn("w:pPr"))
            if pPr is None:
                pPr = OxmlElement("w:pPr")
                p._element.insert(0, pPr)
            old_tabs = pPr.find(qn("w:tabs"))
            if old_tabs is not None:
                pPr.remove(old_tabs)
            tabs = OxmlElement("w:tabs")
            t_r = OxmlElement("w:tab")
            t_r.set(qn("w:val"), "right")
            t_r.set(qn("w:pos"), str(right_pos))
            tabs.append(t_r)
            pPr.append(tabs)
            # 段落左对齐（让 tab 的右对齐 stop 生效）
            jc = pPr.find(qn("w:jc"))
            if jc is None:
                jc = OxmlElement("w:jc")
                pPr.append(jc)
            jc.set(qn("w:val"), "left")

            break  # 命中一条规则即跳
