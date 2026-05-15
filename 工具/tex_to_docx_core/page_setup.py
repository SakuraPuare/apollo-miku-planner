from __future__ import annotations

from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.opc.constants import CONTENT_TYPE as CT, RELATIONSHIP_TYPE as RT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn
from docx.parts.hdrftr import FooterPart


_W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
_R_NS = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"


def _set_page_margins_a4(doc) -> None:
    """A4 幅面 + 上下 2.54cm / 左右 3.17cm（规范 V1/V2/V3）。"""
    from docx.shared import Cm

    for s in doc.sections:
        s.page_width = Cm(21.0)
        s.page_height = Cm(29.7)
        s.top_margin = Cm(2.54)
        s.bottom_margin = Cm(2.54)
        s.left_margin = Cm(3.17)
        s.right_margin = Cm(3.17)


def setup_page_numbers_and_sections(doc) -> None:
    """规范 P1-P5 页码分层：

        Section 1（封面+声明/授权）    ：无页码
        Section 2（目录）              ：罗马数字 I II III...，start=I
        Section 3（摘要）              ：罗马数字 I II III...，start=I
        Section 4（正文"1 绪论"起）     ：阿拉伯数字 1 2 3...，start=1

    每节独立 footer part（FooterPart），避免跨节格式污染。
    """
    _enable_update_fields_on_open(doc)

    body = doc.element.body
    paras = body.findall(qn("w:p"))

    def _para_text(p) -> str:
        return "".join((t.text or "") for t in p.iter(qn("w:t")))

    # 分节锚点定位
    cover_end_idx = None      # 封面+声明/授权 最后一段
    toc_start_idx = None      # 目录第一段
    abstract_start_idx = None # 摘要第一段（论文题目页或"摘要："段）
    body_start_idx = None     # 正文第一段

    for i, p in enumerate(paras):
        t = _para_text(p).strip()
        if cover_end_idx is None and len(t) < 60:
            if "原创性声明" in t or "独创性声明" in t:
                cover_end_idx = i - 1 if i > 0 else None
        if toc_start_idx is None and ("目" in t and "录" in t and len(t) < 10):
            toc_start_idx = i
        if body_start_idx is None:
            pPr = p.find(qn("w:pPr"))
            if pPr is not None:
                pStyle = pPr.find(qn("w:pStyle"))
                if pStyle is not None and pStyle.get(qn("w:val")) == "Heading1":
                    if t.startswith("1 ") or t.startswith("1　") or t.startswith("1绪"):
                        body_start_idx = i

    # 摘要锚点：正文前第一个含"摘要"的段（论文题目页段或摘要正文段）
    if body_start_idx is not None:
        for i in range(body_start_idx - 1, -1, -1):
            t = _para_text(paras[i]).strip()
            if t.startswith("摘要") or "摘要" in t and len(t) < 10:
                abstract_start_idx = i
                break
        # 如果没找到"摘要"标题，找论文题目页（通常在摘要前一两段）
        if abstract_start_idx is None:
            for i in range(body_start_idx - 1, -1, -1):
                t = _para_text(paras[i]).strip()
                if len(t) > 5 and "论文" not in t and "目" not in t:
                    # 第一个有实质内容的段（论文中文题目）
                    abstract_start_idx = i
                    break

    if (
        cover_end_idx is None
        or toc_start_idx is None
        or body_start_idx is None
        or body_start_idx <= cover_end_idx + 1
    ):
        _fallback_single_section(doc)
        return

    # 确定各节结束段
    # 封面节结束：目录开始前一段
    sec1_end_idx = toc_start_idx - 1
    # 目录节结束：摘要开始前一段（如果摘要在目录后）或正文前一段
    if abstract_start_idx is not None and abstract_start_idx > toc_start_idx:
        sec2_end_idx = abstract_start_idx - 1
        sec3_end_idx = body_start_idx - 1
        n_sections = 4
    else:
        # 摘要在目录前或未找到，退化为 3 节
        sec2_end_idx = body_start_idx - 1
        sec3_end_idx = None
        n_sections = 3

    # 清理所有段内旧 sectPr
    for p in paras:
        pPr = p.find(qn("w:pPr"))
        if pPr is None:
            continue
        for sp in list(pPr.findall(qn("w:sectPr"))):
            pPr.remove(sp)

    # 提取模板 sectPr
    template_sectPr = None
    for el in reversed(list(body)):
        if el.tag == qn("w:sectPr"):
            template_sectPr = el
            break
    if template_sectPr is None:
        _fallback_single_section(doc)
        return

    if n_sections == 4:
        # 4 个 FooterPart
        cover_rid = _add_footer_part(doc, _footer_empty_xml())
        toc_rid = _add_footer_part(doc, _footer_page_field_xml())
        abstract_rid = _add_footer_part(doc, _footer_page_field_xml())
        arabic_rid = _add_footer_part(doc, _footer_page_field_xml())

        sec1 = _build_sectPr(template_sectPr, cover_rid, fmt=None, start=None, is_final=False)
        sec2 = _build_sectPr(template_sectPr, toc_rid, fmt="upperRoman", start=1, is_final=False)
        sec3 = _build_sectPr(template_sectPr, abstract_rid, fmt="upperRoman", start=1, is_final=False)
        sec4 = _build_sectPr(template_sectPr, arabic_rid, fmt="decimal", start=1, is_final=True)

        _attach_sectPr_to_para(paras[sec1_end_idx], sec1)
        _attach_sectPr_to_para(paras[sec2_end_idx], sec2)
        _attach_sectPr_to_para(paras[sec3_end_idx], sec3)
        body.remove(template_sectPr)
        body.append(sec4)
    else:
        # 3 节退化
        cover_rid = _add_footer_part(doc, _footer_empty_xml())
        roman_rid = _add_footer_part(doc, _footer_page_field_xml())
        arabic_rid = _add_footer_part(doc, _footer_page_field_xml())

        sec1 = _build_sectPr(template_sectPr, cover_rid, fmt=None, start=None, is_final=False)
        sec2 = _build_sectPr(template_sectPr, roman_rid, fmt="upperRoman", start=1, is_final=False)
        sec3 = _build_sectPr(template_sectPr, arabic_rid, fmt="decimal", start=1, is_final=True)

        _attach_sectPr_to_para(paras[sec1_end_idx], sec1)
        _attach_sectPr_to_para(paras[sec2_end_idx], sec2)
        body.remove(template_sectPr)
        body.append(sec3)


# ---------------------------------------------------------------
# Footer XML 生成
# ---------------------------------------------------------------


def _footer_empty_xml() -> bytes:
    """封面节的空 footer：仅一个空段，不显示任何页码。"""
    return (
        b'<?xml version="1.0" encoding="UTF-8" standalone="yes"?>\n'
        b'<w:ftr xmlns:w="' + _W_NS.encode() + b'" '
        b'xmlns:r="' + _R_NS.encode() + b'">'
        b'<w:p><w:pPr><w:jc w:val="center"/></w:pPr></w:p>'
        b'</w:ftr>'
    )


def _footer_page_field_xml() -> bytes:
    """带 PAGE 域的 footer：居中 + TNR 五号。

    数字格式（decimal / upperRoman）由所属节的 pgNumType 决定，footer XML 自身
    format-agnostic，罗马与阿拉伯两节共享同一模板（只是各自独立 Part）。
    """
    run_rpr = (
        b'<w:rPr>'
        b'<w:rFonts w:ascii="Times New Roman" w:hAnsi="Times New Roman" '
        b'w:eastAsia="Times New Roman" w:cs="Times New Roman"/>'
        b'<w:sz w:val="21"/><w:szCs w:val="21"/>'
        b'</w:rPr>'
    )
    return (
        b'<?xml version="1.0" encoding="UTF-8" standalone="yes"?>\n'
        b'<w:ftr xmlns:w="' + _W_NS.encode() + b'" '
        b'xmlns:r="' + _R_NS.encode() + b'">'
        b'<w:p><w:pPr><w:jc w:val="center"/></w:pPr>'
        b'<w:r>' + run_rpr + b'<w:fldChar w:fldCharType="begin"/></w:r>'
        b'<w:r>' + run_rpr + b'<w:instrText xml:space="preserve"> PAGE \\* MERGEFORMAT </w:instrText></w:r>'
        b'<w:r>' + run_rpr + b'<w:fldChar w:fldCharType="separate"/></w:r>'
        b'<w:r>' + run_rpr + b'<w:t>1</w:t></w:r>'
        b'<w:r>' + run_rpr + b'<w:fldChar w:fldCharType="end"/></w:r>'
        b'</w:p></w:ftr>'
    )


def _add_footer_part(doc, xml_bytes: bytes) -> str:
    """创建 FooterPart 并立即 relate_to 注册，返回 rId。

    FooterPart 第 3 参数必须是 parsed element（非 bytes），否则下游 .paragraphs
    之类的访问会炸（bytes.p_lst 不存在）。
    relation target 必须挂到 *document* part（DocumentPart.relate_to），不是
    package.relate_to —— sectPr 里的 r:id 是基于 document part 的 rels 解析的。
    """
    package = doc.part.package
    partname = package.next_partname("/word/footer%d.xml")
    element = parse_xml(xml_bytes)
    footer_part = FooterPart(partname, CT.WML_FOOTER, element, package)
    return doc.part.relate_to(footer_part, RT.FOOTER)


def _build_sectPr(template, footer_rid: str, *, fmt: str | None,
                  start: int | None, is_final: bool):
    """基于 template sectPr 克隆出新节的 sectPr。

    子元素必须按 OOXML schema 顺序排列，否则 LibreOffice/Word 对 pgNumType
    会静默忽略（实测 upperRoman 被当默认 decimal 渲染）。schema 顺序（精简）：
        headerReference* / footerReference* / footnotePr? / endnotePr? /
        type? / pgSz? / pgMar? / paperSrc? / pgBorders? / lnNumType? /
        pgNumType? / cols? / formProt? / vAlign? / noEndnote? / titlePg? /
        textDirection? / bidi? / rtlGutter? / docGrid? / printerSettings?
    """
    from copy import deepcopy

    sp = deepcopy(template)
    # 清掉旧 footerReference / pgNumType / type（会重建）
    for tag in ("w:footerReference", "w:pgNumType", "w:type"):
        for el in sp.findall(qn(tag)):
            sp.remove(el)

    # 准备要插入的新元素
    new_elems = []
    fref = OxmlElement("w:footerReference")
    fref.set(qn("w:type"), "default")
    fref.set(qn("r:id"), footer_rid)
    new_elems.append(("w:footerReference", fref))

    # 注：不使用 <w:type w:val="continuous"/>。
    # 根因（已实测）：LibreOffice 26.x 对 continuous 节会忽略节内 pgNumType
    # 的 fmt/start，导致 upperRoman / start=1 失效。改用默认 nextPage 语义
    # （不写 w:type），每节自动换新页开始，pgNumType 正常生效。
    # nextPage 的副作用是"换节=换页"，但对本文档恰好是期望行为：
    # 封面 → 换页 → 前置（罗马 I 起）→ 换页 → 正文（阿拉伯 1 起）。

    if fmt or start is not None:
        pgt = OxmlElement("w:pgNumType")
        if fmt:
            pgt.set(qn("w:fmt"), fmt)
        if start is not None:
            pgt.set(qn("w:start"), str(start))
        new_elems.append(("w:pgNumType", pgt))

    # 插入点映射：tag → 插入到首个 "key > tag 顺序权重" 的 child 之前
    _SECTPR_ORDER = [
        "w:headerReference", "w:footerReference",
        "w:footnotePr", "w:endnotePr",
        "w:type",
        "w:pgSz", "w:pgMar", "w:paperSrc", "w:pgBorders",
        "w:lnNumType", "w:pgNumType",
        "w:cols", "w:formProt", "w:vAlign",
        "w:noEndnote", "w:titlePg", "w:textDirection",
        "w:bidi", "w:rtlGutter", "w:docGrid", "w:printerSettings",
    ]
    order_idx = {t: i for i, t in enumerate(_SECTPR_ORDER)}

    def _tag_local(e) -> str:
        # 去掉命名空间前缀，恢复 "w:xxx" 形式
        for prefix_tag in _SECTPR_ORDER:
            if e.tag == qn(prefix_tag):
                return prefix_tag
        return ""

    for tag, new_el in new_elems:
        my_rank = order_idx[tag]
        # 找 sp 中第一个 rank > my_rank 的 child，insert 到它之前
        insert_before = None
        for i, child in enumerate(list(sp)):
            ctag = _tag_local(child)
            if ctag and order_idx.get(ctag, 99) > my_rank:
                insert_before = i
                break
        if insert_before is None:
            sp.append(new_el)
        else:
            sp.insert(insert_before, new_el)
    return sp


def _attach_sectPr_to_para(para, sectPr):
    """把 sectPr 挂到某段 <w:pPr> 内，作为该节的 ending paragraph。"""
    pPr = para.find(qn("w:pPr"))
    if pPr is None:
        pPr = OxmlElement("w:pPr")
        para.insert(0, pPr)
    # 清掉本段已有 sectPr
    for sp in list(pPr.findall(qn("w:sectPr"))):
        pPr.remove(sp)
    pPr.append(sectPr)


def _fallback_single_section(doc) -> None:
    """分节锚点未命中时退化：保留模板单节，仅给末节 footer 加 PAGE 域。

    保底行为不致破坏文档（owner 意识），但不具备 3 节分层语义。
    """
    sect = doc.sections[0]
    footer = sect.footer
    footer.is_linked_to_previous = False
    if footer.paragraphs:
        p = footer.paragraphs[0]
        for r in list(p.runs):
            r._element.getparent().remove(r._element)
    else:
        p = footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), " PAGE ")
    inner_r = OxmlElement("w:r")
    inner_rPr = OxmlElement("w:rPr")
    rF = OxmlElement("w:rFonts")
    rF.set(qn("w:ascii"), "Times New Roman")
    rF.set(qn("w:hAnsi"), "Times New Roman")
    rF.set(qn("w:eastAsia"), "Times New Roman")
    inner_rPr.append(rF)
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), "21")
    inner_rPr.append(sz)
    szCs = OxmlElement("w:szCs")
    szCs.set(qn("w:val"), "21")
    inner_rPr.append(szCs)
    inner_r.append(inner_rPr)
    inner_t = OxmlElement("w:t")
    inner_t.text = "1"
    inner_r.append(inner_t)
    fld.append(inner_r)
    p._element.append(fld)


def _enable_update_fields_on_open(doc) -> None:
    """让 Word/WPS 打开文档时更新 TOC/PAGE 等域，减少目录页码不一致。"""
    settings = doc.settings.element
    update_fields = settings.find(qn("w:updateFields"))
    if update_fields is None:
        update_fields = OxmlElement("w:updateFields")
        settings.append(update_fields)
    update_fields.set(qn("w:val"), "true")
