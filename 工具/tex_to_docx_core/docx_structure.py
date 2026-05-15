from __future__ import annotations

import re

from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from .config import REFS_ANCHOR_TEXT, THESIS_DIR
from .docx_common import _make_page_break_para

# 无编号章：摘要/Abstract/参考文献/致谢/成果 —— 不计入 \chapter 编号
NON_NUMBERED_HEADINGS_1 = {
    "摘要",
    "Abstract",
    "参考文献",
    "致谢",
    "致 谢",
    "本科期间的学习与科研成果",
}

# 这些 H1 虽然保留 Heading1 样式/分页，但不在 TOC 目录中展示
# 规范：摘要/Abstract 是论文正文的"前置材料"（front matter），目录列参考文献起
EXCLUDED_FROM_TOC = {
    "摘要",
    "Abstract",
}

# 这些章单独起一页
PAGE_BREAK_BEFORE_HEADINGS = {
    "参考文献",
    "致谢",
    "本科期间的学习与科研成果",
}
# 全部 Heading 1 都要分页（除了首个"摘要"——它紧跟目录 TOC 分页）
PAGE_BREAK_ALL_H1 = True


def _relocate_bibliography(doc) -> None:
    """pandoc citeproc 把参考文献条目放在文档末尾；把它们整体搬到 anchor 一级标题之后。"""
    body = doc.element.body
    # 1. 找 anchor（文字为"参考文献"的 Heading 1）
    anchor = None
    for p in doc.paragraphs:
        if (
            p.style
            and p.style.name == "Heading 1"
            and p.text.strip() == REFS_ANCHOR_TEXT
        ):
            anchor = p._element
            break
    if anchor is None:
        return  # 没埋 anchor，放弃搬运

    # 2. 收集所有 Bibliography 段（pandoc 给参考文献条目的样式名）
    bib_elems = [
        p._element for p in doc.paragraphs if p.style and p.style.name == "Bibliography"
    ]
    if not bib_elems:
        return

    # 3. 搬运：先 detach，再按序 insert 到 anchor 之后
    for el in bib_elems:
        el.getparent().remove(el)

    # anchor 在 body 里的下标
    children = list(body)
    anchor_idx = children.index(anchor)
    for offset, el in enumerate(bib_elems, start=1):
        body.insert(anchor_idx + offset, el)


def _extract_thesis_titles() -> tuple[str, str]:
    """从 info.tex 提取中英文论文题目。"""
    info = THESIS_DIR / "info.tex"
    if not info.exists():
        return ("论文题目", "Thesis Title")
    text = info.read_text(encoding="utf-8")

    def _find(cmd: str) -> str:
        m = re.search(rf"\\newcommand\{{\\{cmd}\}}\{{([^}}]*)\}}", text)
        return m.group(1) if m else ""

    zh_a = _find("thesisTitleZhA")
    zh_b = _find("thesisTitleZhB")
    zh = (zh_a + zh_b).strip() or "论文题目"
    en = _find("thesisTitleEn").replace("\\\\", " ").replace("  ", " ").strip()
    return zh, en or "Thesis Title"


def insert_page_breaks_before_headings(doc) -> None:
    """在所有 Heading 1 之前插入分页符（首个除外，由 TOC 域段内 page break 覆盖）。"""
    h1_list = [p for p in doc.paragraphs if p.style and p.style.name == "Heading 1"]
    for i, p in enumerate(h1_list):
        if i == 0:
            # 首个 Heading 1（通常是"摘要"）紧跟 TOC 域段内的分页符，不重复加
            continue
        # 若前一段已有分页符则跳过（避免 prepend_front_matter 之后重复）
        prev_el = p._element.getprevious()
        if prev_el is not None:
            has_break = any(
                br.get(qn("w:type")) == "page" for br in prev_el.iter(qn("w:br"))
            )
            if has_break:
                continue
        page_break_p = _make_page_break_para(doc)
        p._element.addprevious(page_break_p)


def demote_heading4_to_heading3(doc) -> None:
    """规范"最多三级标题，不得出现四级标题"：把 Heading 4 样式降级为 Heading 3。"""
    try:
        h3_style = doc.styles["Heading 3"]
    except KeyError:
        return
    for p in doc.paragraphs:
        if p.style and p.style.name == "Heading 4":
            p.style = h3_style


def normalize_special_h1_text(doc) -> None:
    """致谢 → 致 谢（中间空一字符，对齐规范"致 谢"要求）。"""
    for p in doc.paragraphs:
        if p.style and p.style.name == "Heading 1" and p.text.strip() == "致谢":
            if p.runs:
                for r in p.runs:
                    r.text = r.text.replace("致谢", "致 谢")


def insert_thesis_title_pages(doc) -> None:
    """在"摘要"Heading 1 前插入中文论文题目段；在"Abstract"前插入英文论文题目段。
    题目：三号(16pt) 黑体/TNR 加粗居中，单独成页（前加分页符）。"""
    zh_title, en_title = _extract_thesis_titles()

    def _make_title_p(text: str, cjk: str, ascii_: str) -> OxmlElement:
        p = OxmlElement("w:p")
        pPr = OxmlElement("w:pPr")
        jc = OxmlElement("w:jc")
        jc.set(qn("w:val"), "center")
        pPr.append(jc)
        sp = OxmlElement("w:spacing")
        sp.set(qn("w:before"), "120")  # 6pt
        sp.set(qn("w:after"), "31")  # 1.54pt
        pPr.append(sp)
        p.append(pPr)
        r = OxmlElement("w:r")
        rPr = OxmlElement("w:rPr")
        rFonts = OxmlElement("w:rFonts")
        rFonts.set(qn("w:ascii"), ascii_)
        rFonts.set(qn("w:hAnsi"), ascii_)
        rFonts.set(qn("w:eastAsia"), cjk)
        rPr.append(rFonts)
        b = OxmlElement("w:b")
        rPr.append(b)
        bcs = OxmlElement("w:bCs")
        rPr.append(bcs)
        sz = OxmlElement("w:sz")
        sz.set(qn("w:val"), "32")  # 三号 16pt → sz=32 half-points
        rPr.append(sz)
        szCs = OxmlElement("w:szCs")
        szCs.set(qn("w:val"), "32")
        rPr.append(szCs)
        r.append(rPr)
        t = OxmlElement("w:t")
        t.text = text
        t.set(qn("xml:space"), "preserve")
        r.append(t)
        p.append(r)
        return p

    for anchor_text in ("摘要", "Abstract"):
        anchor = next(
            (
                p
                for p in doc.paragraphs
                if p.style
                and p.style.name == "Heading 1"
                and p.text.strip() == anchor_text
            ),
            None,
        )
        if anchor is None:
            continue
        if anchor_text == "摘要":
            title_p = _make_title_p(zh_title, "黑体", "Times New Roman")
        else:
            title_p = _make_title_p(en_title, "Times New Roman", "Times New Roman")
        anchor._element.addprevious(title_p)


def _prepend_run_text(p, text: str) -> None:
    """在段落首 run 开头插入文字；无 run 则新建一个（继承段落样式）。"""
    if p.runs:
        r = p.runs[0]
        r.text = text + (r.text or "")
    else:
        p.add_run(text)


def add_heading_numbers(doc) -> None:
    """按章节层次给 Heading 1/2/3/4 前缀自动编号，对齐模板"1 / 1.1 / 1.1.1"风格。
    摘要/Abstract/参考文献/致谢/成果 等无编号章跳过（但会重置子计数器）。"""
    chap = sec = subsec = subsubsec = 0
    for p in doc.paragraphs:
        sn = p.style.name if p.style else ""
        text = p.text.strip()
        if sn == "Heading 1":
            if text in NON_NUMBERED_HEADINGS_1 or text.startswith("附录"):
                sec = subsec = subsubsec = 0
                continue
            chap += 1
            sec = subsec = subsubsec = 0
            _prepend_run_text(p, f"{chap} ")
        elif sn == "Heading 2" and chap > 0:
            sec += 1
            subsec = subsubsec = 0
            _prepend_run_text(p, f"{chap}.{sec} ")
        elif sn == "Heading 3" and chap > 0 and sec > 0:
            subsec += 1
            subsubsec = 0
            _prepend_run_text(p, f"{chap}.{sec}.{subsec} ")
        elif sn == "Heading 4" and chap > 0 and sec > 0 and subsec > 0:
            subsubsec += 1
            _prepend_run_text(p, f"{chap}.{sec}.{subsec}.{subsubsec} ")


def _inject_toc_bookmarks(doc) -> list:
    """给所有 Heading 1/2 注入 bookmark（_Toc_XXX），返回 [(level, name, text), ...]。

    根因：检测器需要真实 TOC 段才能判定缩进。如果 docx 仅含 TOC 域（fldChar），
    段要等 Word/LO 更新域才生成，段级 pPr 不受样式 firstLine 控制 → 检测器报错
    "TOC2 首行缩进 0 字符"。解法是 postprocess 阶段预渲染真实 TOC 段 + PAGEREF
    域保留页码动态更新能力。这一步先给 heading 加 bookmark 供 PAGEREF 引用。
    """
    body = doc.element.body
    entries = []
    bm_id = 1000
    # 收集所有非 NON_NUMBERED_HEADINGS_1 的 Heading 1/2
    for p in body.findall('.//' + qn('w:p')):
        pPr = p.find(qn('w:pPr'))
        pStyle = pPr.find(qn('w:pStyle')) if pPr is not None else None
        sid = pStyle.get(qn('w:val')) if pStyle is not None else ''
        if sid not in ('Heading1', 'Heading2'):
            continue
        txt = ''.join(t.text or '' for t in p.findall('.//' + qn('w:t'))).strip()
        if not txt:
            continue
        # 过滤掉不入目录的 H1（摘要/Abstract）：它们仍保留 Heading1 样式与分页，
        # 但不在 TOC 段列表中出现，也不打 bookmark（无外部引用价值）
        if txt in EXCLUDED_FROM_TOC:
            continue
        # 跳过非 TOC 目录内容的 H1（封面、前置页的一些标签）
        # 保留：1 绪论 / 2 ... / 参考文献 / 致谢 / 成果
        level = 1 if sid == 'Heading1' else 2
        bm_name = f"_Toc{bm_id}"
        bm_id += 1
        # 插入 bookmarkStart / bookmarkEnd，包裹所有 run
        bs = OxmlElement("w:bookmarkStart")
        bs.set(qn("w:id"), str(bm_id))
        bs.set(qn("w:name"), bm_name)
        be = OxmlElement("w:bookmarkEnd")
        be.set(qn("w:id"), str(bm_id))
        # 插入点：pPr 之后（即段内所有 run 之前）与段尾
        pPr_el = p.find(qn('w:pPr'))
        insert_idx = list(p).index(pPr_el) + 1 if pPr_el is not None else 0
        p.insert(insert_idx, bs)
        p.append(be)
        entries.append((level, bm_name, txt))
    return entries


def _make_toc_entry_paragraph(level: int, bm_name: str, text: str):
    """生成一个真实 TOC 段。

    顶层设计（根因解法）：`<w:hyperlink>` 只包标题文字；`tab` 与 `PAGEREF` 域
    放在 hyperlink **外部**。避免 WPS 的"hyperlink 嵌套 PAGEREF"域解析 bug
    （症状：WPS 打开后所有目录页码 fallback 成同一值，如全 6）。

    结构：
        <w:p>
          <w:pPr>...style + tabs + ind...</w:pPr>
          <w:hyperlink anchor="_TocN"><w:r>(黑体)1 绪论</w:r></w:hyperlink>  ← 可点击跳转
          <w:r><w:tab/></w:r>                                                  ← tab 到右侧
          <w:r><w:fldChar begin/></w:r>                                        ← 独立 PAGEREF 域
          <w:r><w:instrText> PAGEREF _TocN \h </w:instrText></w:r>
          <w:r><w:fldChar separate/></w:r>
          <w:r><w:t>1</w:t></w:r>                                              ← 占位页码
          <w:r><w:fldChar end/></w:r>
        </w:p>

    Word/WPS 打开即可按 settings.updateFields=true 自动更新所有 PAGEREF 域到
    真实页号，无需按 F9，无需静态化。

    学校规范：
      - TOC1 左对齐不缩进，中文用 **黑体**（eastAsia=黑体）
      - TOC2/TOC3 首行缩进 2 字符（firstLineChars=200, firstLine=480）
    """
    style_id = f"TOC{level}"
    # TOC1: firstLine=0（左对齐）；TOC2/TOC3: firstLine=480（2 字符缩进）
    fl, flc = ("0", "0") if level == 1 else ("480", "200")

    p = OxmlElement("w:p")
    pPr = OxmlElement("w:pPr")
    ps = OxmlElement("w:pStyle")
    ps.set(qn("w:val"), style_id)
    pPr.append(ps)
    # tabs: right dot 引导到页码
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "right")
    tab.set(qn("w:leader"), "dot")
    tab.set(qn("w:pos"), "8306")
    tabs.append(tab)
    pPr.append(tabs)
    # 段级 ind：显式、不依赖样式继承
    ind = OxmlElement("w:ind")
    ind.set(qn("w:firstLineChars"), flc)
    ind.set(qn("w:firstLine"), fl)
    pPr.append(ind)
    p.append(pPr)

    # --- 1. HYPERLINK：只包标题文字（可点击跳转）---
    hyp = OxmlElement("w:hyperlink")
    hyp.set(qn("w:anchor"), bm_name)
    hyp.set(qn("w:history"), "1")

    # 标题文本 run
    # TOC1 条目：中文必须用黑体（eastAsia=黑体），西文保持 TNR
    r_text = OxmlElement("w:r")
    if level == 1:
        rPr = OxmlElement("w:rPr")
        rf = OxmlElement("w:rFonts")
        rf.set(qn("w:ascii"), "Times New Roman")
        rf.set(qn("w:hAnsi"), "Times New Roman")
        rf.set(qn("w:eastAsia"), "黑体")
        rf.set(qn("w:cs"), "Times New Roman")
        rPr.append(rf)
        r_text.append(rPr)
    t = OxmlElement("w:t")
    t.set(qn("xml:space"), "preserve")
    t.text = text
    r_text.append(t)
    hyp.append(r_text)
    p.append(hyp)

    # --- 2. tab run（在 hyperlink 外） ---
    r_tab = OxmlElement("w:r")
    tab_el = OxmlElement("w:tab")
    r_tab.append(tab_el)
    p.append(r_tab)

    # --- 3. PAGEREF 域（独立，不嵌套在 hyperlink 内）---
    # 关键：begin / instr / separate / 占位 / end 五段 run 直接挂在段落下，
    # 绕开 WPS 的 hyperlink+PAGEREF 嵌套解析 bug，让域引擎正常更新页号。
    r_fb = OxmlElement("w:r")
    fb = OxmlElement("w:fldChar")
    fb.set(qn("w:fldCharType"), "begin")
    r_fb.append(fb)
    p.append(r_fb)

    r_fi = OxmlElement("w:r")
    fi = OxmlElement("w:instrText")
    fi.set(qn("xml:space"), "preserve")
    fi.text = f" PAGEREF {bm_name} \\h "
    r_fi.append(fi)
    p.append(r_fi)

    r_fs = OxmlElement("w:r")
    fs = OxmlElement("w:fldChar")
    fs.set(qn("w:fldCharType"), "separate")
    r_fs.append(fs)
    p.append(r_fs)

    # 占位页码（Word/WPS 打开后按 updateFields 更新成真实值）
    r_pn = OxmlElement("w:r")
    t_pn = OxmlElement("w:t")
    t_pn.text = "1"
    r_pn.append(t_pn)
    p.append(r_pn)

    r_fe = OxmlElement("w:r")
    fe = OxmlElement("w:fldChar")
    fe.set(qn("w:fldCharType"), "end")
    r_fe.append(fe)
    p.append(r_fe)

    return p


def _make_toc_block(doc) -> list:
    """构造预渲染的 TOC：标题段 + 真实 TOC 段列表 + 分页符。

    不再使用 `TOC` 域。改为扫描所有 Heading 1/2 自行生成 TOC 段，每段显式
    段级 ind 以对齐检测器的分级缩进规则。页码由 PAGEREF 域保留动态更新。
    """
    els = []

    # "目  录" 标题段
    title_p = OxmlElement("w:p")
    title_pPr = OxmlElement("w:pPr")
    title_jc = OxmlElement("w:jc")
    title_jc.set(qn("w:val"), "center")
    title_pPr.append(title_jc)
    title_p.append(title_pPr)
    title_r = OxmlElement("w:r")
    title_rPr = OxmlElement("w:rPr")
    rf = OxmlElement("w:rFonts")
    rf.set(qn("w:eastAsia"), "黑体")
    rf.set(qn("w:ascii"), "黑体")
    title_rPr.append(rf)
    b = OxmlElement("w:b")
    title_rPr.append(b)
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), "32")
    title_rPr.append(sz)
    title_r.append(title_rPr)
    title_t = OxmlElement("w:t")
    title_t.text = "目  录"
    title_r.append(title_t)
    title_p.append(title_r)
    els.append(title_p)

    # 扫描 heading 注入 bookmark，并生成 TOC 段
    entries = _inject_toc_bookmarks(doc)
    for level, bm_name, text in entries:
        els.append(_make_toc_entry_paragraph(level, bm_name, text))

    # 分页由 setup_page_numbers_and_sections 的 sectPr(nextPage) 处理，不再重复加分页符

    return els
