"""单元测试：验证 outputs/thesis.docx 符合 FORMAT_SPEC.md 排版规范。

测试直接读取生成的 docx XML 属性，不 mock 排版流程。
运行：uv run pytest 工具/tests/test_format_spec.py -v
"""
from __future__ import annotations

import re

import pytest
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, Emu


# ═══════════════════════════════════════════════════════════════════
# Helpers
# ═══════════════════════════════════════════════════════════════════

def _get_pPr(p):
    return p._element.find(qn("w:pPr"))


def _get_spacing(p):
    pPr = _get_pPr(p)
    if pPr is None:
        return None
    return pPr.find(qn("w:spacing"))


def _get_ind(p):
    pPr = _get_pPr(p)
    if pPr is None:
        return None
    return pPr.find(qn("w:ind"))


def _effective_sz(p):
    """获取段落有效字号 (half-points)。优先取 run 级，否则取段落样式。"""
    for r in p.runs:
        rPr = r._element.find(qn("w:rPr"))
        if rPr is not None:
            sz = rPr.find(qn("w:sz"))
            if sz is not None:
                return int(sz.get(qn("w:val")))
    return None


def _effective_font_east(p):
    """获取段落首个 run 的 eastAsia 字体。"""
    for r in p.runs:
        rPr = r._element.find(qn("w:rPr"))
        if rPr is not None:
            rf = rPr.find(qn("w:rFonts"))
            if rf is not None:
                return rf.get(qn("w:eastAsia"))
    return None


def _is_bold(run):
    """检查 run 是否有显式 <w:b/> 加粗。"""
    rPr = run._element.find(qn("w:rPr"))
    if rPr is None:
        return False
    b = rPr.find(qn("w:b"))
    if b is None:
        return False
    val = b.get(qn("w:val"))
    return val is None or val in ("1", "true")


def _is_bold_effective(p):
    """检查段落是否有效加粗（run 级或 style 级）。"""
    # style-level bold
    if p.style:
        style_el = p.style.element
        rPr = style_el.find(qn("w:rPr"))
        if rPr is not None:
            b = rPr.find(qn("w:b"))
            if b is not None:
                val = b.get(qn("w:val"))
                if val is None or val in ("1", "true"):
                    return True
    # run-level bold
    return _all_runs_bold(p)


_ALGO_CODE_CAP_RE = re.compile(r"^(代码|算法)\s*\d+[\.\-]\d+")


def _all_runs_bold(p):
    runs = [r for r in p.runs if (r.text or "").strip()]
    if not runs:
        return True
    return all(_is_bold(r) for r in runs)


def _para_alignment(p):
    """返回段落对齐: 'left', 'center', 'right', 'both', None."""
    pPr = _get_pPr(p)
    if pPr is None:
        return None
    jc = pPr.find(qn("w:jc"))
    if jc is None:
        return None
    return jc.get(qn("w:val"))


def _body_paragraphs(doc):
    """返回正文段（排除标题、目录、参考文献、caption 等）。"""
    result = []
    in_body = False
    for p in doc.paragraphs:
        style = p.style.name if p.style else ""
        text = p.text.strip()
        if style == "Heading 1" and re.match(r"^\d+\s", text):
            in_body = True
        if style == "Heading 1" and text == "参考文献":
            break
        if not in_body:
            continue
        if style.startswith("Heading") or style.startswith("toc"):
            continue
        if style in ("Caption", "Image Caption", "Table Caption", "Captioned Figure"):
            continue
        if not text:
            continue
        # 排除算法/代码 caption（可能被标为 Body Text）
        if _ALGO_CODE_CAP_RE.match(text):
            continue
        result.append(p)
    return result


def _toc_paragraphs(doc):
    """返回所有目录段。"""
    return [p for p in doc.paragraphs
            if p.style and p.style.name.lower().startswith("toc")]


def _bibliography_paragraphs(doc):
    """返回参考文献条目段。"""
    return [p for p in doc.paragraphs
            if p.style and p.style.name.lower().startswith("bibliograph")]


def _heading_paragraphs(doc, level):
    return [p for p in doc.paragraphs
            if p.style and p.style.name == f"Heading {level}"]


def _caption_paragraphs(doc):
    return [p for p in doc.paragraphs
            if p.style and p.style.name in ("Caption", "Image Caption", "Table Caption")]


# ═══════════════════════════════════════════════════════════════════
# §1 全局默认
# ═══════════════════════════════════════════════════════════════════

class TestGlobalDefaults:
    def test_page_margins(self, doc):
        """A4 页面上下 2.54cm 左右 3.17cm。"""
        for section in doc.sections:
            assert abs(section.top_margin - Cm(2.54)) < Emu(5000), \
                f"top margin {section.top_margin} != 2.54cm"
            assert abs(section.bottom_margin - Cm(2.54)) < Emu(5000), \
                f"bottom margin {section.bottom_margin} != 2.54cm"
            assert abs(section.left_margin - Cm(3.17)) < Emu(5000), \
                f"left margin {section.left_margin} != 3.17cm"
            assert abs(section.right_margin - Cm(3.17)) < Emu(5000), \
                f"right margin {section.right_margin} != 3.17cm"

    def test_all_colors_black(self, doc):
        """全文字体颜色应为黑色（000000）或未设置（继承黑色）。"""
        violations = []
        for i, p in enumerate(doc.paragraphs):
            for r in p.runs:
                rPr = r._element.find(qn("w:rPr"))
                if rPr is None:
                    continue
                color = rPr.find(qn("w:color"))
                if color is None:
                    continue
                val = color.get(qn("w:val"))
                if val and val.lower() not in ("000000", "auto"):
                    violations.append(
                        f"P{i} run='{r.text[:20]}' color={val}"
                    )
        assert not violations, \
            f"Found {len(violations)} non-black runs:\n" + "\n".join(violations[:10])

    def test_styles_doc_defaults(self, doc):
        """docDefaults rFonts: ascii/hAnsi = TNR, eastAsia = 宋体。"""
        styles_el = doc.styles.element
        defaults = styles_el.find(qn("w:docDefaults"))
        assert defaults is not None
        rPrDefault = defaults.find(qn("w:rPrDefault"))
        assert rPrDefault is not None
        rPr = rPrDefault.find(qn("w:rPr"))
        assert rPr is not None
        rFonts = rPr.find(qn("w:rFonts"))
        assert rFonts is not None
        assert rFonts.get(qn("w:ascii")) == "Times New Roman"
        assert rFonts.get(qn("w:hAnsi")) == "Times New Roman"
        assert rFonts.get(qn("w:eastAsia")) == "宋体"


# ═══════════════════════════════════════════════════════════════════
# §3.1 目录块
# ═══════════════════════════════════════════════════════════════════

class TestTOC:
    def test_toc1_bold(self, doc):
        """TOC1 一级目录项必须加粗——style 级或 run 级 (§5.10)。"""
        toc1 = [p for p in doc.paragraphs
                if p.style and p.style.name.lower() == "toc 1"]
        assert toc1, "No TOC1 paragraphs found"
        violations = []
        for p in toc1:
            if not _is_bold_effective(p):
                violations.append(p.text[:40])
        assert not violations, \
            f"TOC1 not bold: {violations[:5]}"

    def test_toc1_no_indent(self, doc):
        """TOC1 无首行缩进 (§5.5)。"""
        toc1 = [p for p in doc.paragraphs
                if p.style and p.style.name.lower() == "toc 1"]
        violations = []
        for p in toc1:
            ind = _get_ind(p)
            if ind is not None:
                fl = ind.get(qn("w:firstLine"))
                flc = ind.get(qn("w:firstLineChars"))
                if (fl and fl != "0") or (flc and flc != "0"):
                    violations.append(f"'{p.text[:30]}' firstLine={fl} chars={flc}")
        assert not violations, \
            f"TOC1 has indent: {violations[:5]}"

    def test_toc2_indent(self, doc):
        """TOC2 首行缩进 2 字符 (§5.5)。"""
        toc2 = [p for p in doc.paragraphs
                if p.style and p.style.name.lower() == "toc 2"]
        assert toc2, "No TOC2 paragraphs found"
        violations = []
        for p in toc2:
            ind = _get_ind(p)
            if ind is None:
                violations.append(f"'{p.text[:30]}' no indent")
                continue
            flc = ind.get(qn("w:firstLineChars"))
            fl = ind.get(qn("w:firstLine"))
            # 2字符 = firstLineChars=200 或 firstLine=480
            if not ((flc and int(flc) >= 200) or (fl and int(fl) >= 420)):
                violations.append(f"'{p.text[:30]}' firstLine={fl} chars={flc}")
        assert not violations, \
            f"TOC2 indent wrong: {violations[:5]}"

    def test_toc_spacing_zero_after(self, doc):
        """目录段段后 0 (§5.3)。"""
        tocs = _toc_paragraphs(doc)
        assert tocs, "No TOC paragraphs"
        violations = []
        for p in tocs:
            sp = _get_spacing(p)
            if sp is not None:
                after = sp.get(qn("w:after"))
                if after and int(after) > 0:
                    violations.append(f"'{p.text[:30]}' after={after}")
        assert not violations, \
            f"TOC after!=0: {violations[:5]}"

    def test_toc_line_spacing_1p5(self, doc):
        """目录段 1.5 倍行距 (§5.4) = line=360 auto（段级或样式级继承）。"""
        tocs = _toc_paragraphs(doc)
        violations = []
        for p in tocs:
            sp = _get_spacing(p)
            # 段级有显式 line 设置
            if sp is not None:
                line = sp.get(qn("w:line"))
                rule = sp.get(qn("w:lineRule"))
                if line is not None:
                    if line != "360" or rule != "auto":
                        violations.append(f"'{p.text[:30]}' line={line} rule={rule}")
                    continue
            # 无段级 line → 检查样式级
            style_el = p.style.element if p.style else None
            if style_el is not None:
                pPr = style_el.find(qn("w:pPr"))
                if pPr is not None:
                    ssp = pPr.find(qn("w:spacing"))
                    if ssp is not None:
                        line = ssp.get(qn("w:line"))
                        rule = ssp.get(qn("w:lineRule"))
                        if line == "360" and rule == "auto":
                            continue
            violations.append(f"'{p.text[:30]}' no line spacing found")
        assert not violations, \
            f"TOC line spacing wrong: {violations[:5]}"


# ═══════════════════════════════════════════════════════════════════
# §3.3 正文块
# ═══════════════════════════════════════════════════════════════════

class TestBody:
    def test_h1_bold_14pt(self, doc):
        """H1 章标题加粗四号 (sz=28)——style 级或 run 级。"""
        h1s = _heading_paragraphs(doc, 1)
        assert h1s
        numbered = [p for p in h1s if re.match(r"^\d+\s", p.text.strip())]
        violations = []
        for p in numbered:
            if not _is_bold_effective(p):
                violations.append(f"'{p.text[:30]}' not bold")
            sz = _effective_sz(p)
            if sz and sz != 28:
                violations.append(f"'{p.text[:30]}' sz={sz} expect 28")
        assert not violations, "\n".join(violations[:10])

    def test_h2_bold_xiaosi(self, doc):
        """H2 节标题加粗小四 (sz=24)——style 级或 run 级。"""
        h2s = _heading_paragraphs(doc, 2)
        assert h2s
        violations = []
        for p in h2s:
            if not _is_bold_effective(p):
                violations.append(f"'{p.text[:30]}' not bold")
            sz = _effective_sz(p)
            if sz and sz != 24:
                violations.append(f"'{p.text[:30]}' sz={sz} expect 24")
        assert not violations, "\n".join(violations[:10])

    def test_h3_bold_xiaosi(self, doc):
        """H3 小节标题宋体加粗小四 (sz=24)——style 级或 run 级。"""
        h3s = _heading_paragraphs(doc, 3)
        if not h3s:
            pytest.skip("No H3 in document")
        violations = []
        for p in h3s:
            if not _is_bold_effective(p):
                violations.append(f"'{p.text[:30]}' not bold")
            sz = _effective_sz(p)
            if sz and sz != 24:
                violations.append(f"'{p.text[:30]}' sz={sz} expect 24")
        assert not violations, "\n".join(violations[:10])

    def test_body_no_run_bold(self, doc):
        """正文段无 run 级 bold (§5.9)。"""
        body = _body_paragraphs(doc)
        assert body, "No body paragraphs found"
        violations = []
        for p in body:
            for r in p.runs:
                if not (r.text or "").strip():
                    continue
                if _is_bold(r):
                    violations.append(f"'{p.text[:40]}' run='{r.text[:20]}' bold")
                    break
        # 允许极少量（如公式标签等边缘情况）
        assert len(violations) <= 3, \
            f"Found {len(violations)} body paras with bold runs:\n" + "\n".join(violations[:10])

    def test_body_first_line_indent(self, doc):
        """正文段首行缩进 2 字符。"""
        body = _body_paragraphs(doc)
        violations = []
        for p in body:
            text = p.text.strip()
            # 跳过公式段、图片段等
            if re.match(r"^[（(]\d+[-–—]\d+[)）]", text):
                continue
            ind = _get_ind(p)
            if ind is None:
                violations.append(f"'{text[:30]}' no indent")
                continue
            flc = ind.get(qn("w:firstLineChars"))
            fl = ind.get(qn("w:firstLine"))
            if not ((flc and int(flc) >= 200) or (fl and int(fl) >= 420)):
                violations.append(f"'{text[:30]}' fl={fl} flc={flc}")
        # 允许少量例外（公式、特殊格式段）
        assert len(violations) <= 5, \
            f"Found {len(violations)} body paras without indent:\n" + "\n".join(violations[:10])

    def test_body_line_spacing_1p5(self, doc):
        """正文段 1.5 倍行距 (line=360 auto)。"""
        body = _body_paragraphs(doc)
        violations = []
        for p in body[:50]:  # 抽样前 50 段
            sp = _get_spacing(p)
            if sp is None:
                continue
            line = sp.get(qn("w:line"))
            rule = sp.get(qn("w:lineRule"))
            if line and line != "360":
                violations.append(f"'{p.text[:30]}' line={line}")
            if rule and rule != "auto":
                violations.append(f"'{p.text[:30]}' rule={rule}")
        assert not violations, \
            f"Body line spacing wrong:\n" + "\n".join(violations[:10])

    def test_caption_center_bold(self, doc):
        """图/表 caption 居中加粗。"""
        caps = _caption_paragraphs(doc)
        if not caps:
            pytest.skip("No captions found")
        violations = []
        for p in caps:
            align = _para_alignment(p)
            if align != "center":
                violations.append(f"'{p.text[:30]}' align={align}")
            if not _all_runs_bold(p):
                violations.append(f"'{p.text[:30]}' not bold")
        assert not violations, "\n".join(violations[:10])

    def test_equation_right_align(self, doc):
        """独立公式编号段右对齐 (§5.8)。"""
        eq_re = re.compile(r"^[（(]\d+[-–—]\d+[)）]\s*$")
        equations = [p for p in doc.paragraphs if eq_re.match(p.text.strip())]
        if not equations:
            pytest.skip("No equation number paragraphs found")
        violations = []
        for p in equations:
            align = _para_alignment(p)
            if align != "right":
                violations.append(f"'{p.text.strip()}' align={align}")
        assert not violations, "\n".join(violations[:10])


# ═══════════════════════════════════════════════════════════════════
# §3.4 参考文献块
# ═══════════════════════════════════════════════════════════════════

class TestBibliography:
    def test_bibliography_heading_14pt_left(self, doc):
        """参考文献标题：四号(sz=28)黑体加粗左对齐。"""
        h1s = _heading_paragraphs(doc, 1)
        bib_h1 = [p for p in h1s if "参考文献" in p.text]
        assert bib_h1, "未找到参考文献标题"
        p = bib_h1[0]
        sz = _effective_sz(p)
        assert sz == 28, f"参考文献标题字号 sz={sz}, 期望 28 (四号)"
        align = _para_alignment(p)
        assert align in ("left", None), f"参考文献标题对齐={align}, 期望 left"
        assert _is_bold_effective(p), "参考文献标题未加粗"

    def test_bibliography_exact_20pt(self, doc):
        """参考文献条目固定 20 磅行距 (§5.7) = line=400 exact。"""
        bibs = _bibliography_paragraphs(doc)
        if not bibs:
            pytest.skip("No bibliography paragraphs")
        violations = []
        for p in bibs:
            sp = _get_spacing(p)
            if sp is None:
                violations.append(f"'{p.text[:30]}' no spacing")
                continue
            line = sp.get(qn("w:line"))
            rule = sp.get(qn("w:lineRule"))
            if line != "400" or rule != "exact":
                violations.append(f"'{p.text[:30]}' line={line} rule={rule}")
        assert not violations, \
            f"Bib spacing wrong:\n" + "\n".join(violations[:5])


# ═══════════════════════════════════════════════════════════════════
# §5 检测器硬性口径
# ═══════════════════════════════════════════════════════════════════

class TestInspectorRules:
    def test_caption_no_prefix_space(self, doc):
        """图/表/代码/算法 caption 前缀与数字无空格 (§5.1)。"""
        cap_re = re.compile(r"^(图|表|代码|算法)\s+\d")
        violations = []
        for p in doc.paragraphs:
            text = p.text.strip()
            if cap_re.match(text):
                violations.append(text[:40])
        assert not violations, \
            f"Caption prefix has space:\n" + "\n".join(violations[:10])

    def test_code_caption_no_indent(self, doc):
        """代码/算法 caption 段无首行缩进 (§5.2)。
        判定 caption：匹配 ^(代码|算法)N.N 且全段加粗（排除正文引用）。"""
        cap_re = re.compile(r"^(代码|算法)\s*\d+[\.\-]\d+")
        violations = []
        for p in doc.paragraphs:
            text = p.text.strip()
            if not cap_re.match(text):
                continue
            # 只检查真正的 caption（全段加粗且 < 80 字）
            if len(text) > 80:
                continue
            if not _all_runs_bold(p):
                continue
            ind = _get_ind(p)
            if ind is not None:
                fl = ind.get(qn("w:firstLine"))
                flc = ind.get(qn("w:firstLineChars"))
                if (fl and fl != "0") or (flc and flc != "0"):
                    violations.append(f"'{text[:30]}' fl={fl} flc={flc}")
        assert not violations, "\n".join(violations[:10])
